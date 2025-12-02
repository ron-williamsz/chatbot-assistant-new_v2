from flask import Blueprint, request, jsonify, render_template, current_app, redirect, url_for, make_response, g, send_file, Response
from app.services.openai_client import OpenAIAssistantClient
import requests
import time
from app import database as db
from functools import wraps
import os
import uuid
from datetime import datetime
from threading import Thread
from .services.openai_service import OpenAIService
from .services.database import get_assistants, save_assistant, mark_assistant_deleted
from .services.transcriber_client import transcriber_client
import json
import subprocess
import shutil
import zipfile
from werkzeug.utils import secure_filename
import math

bp = Blueprint('main', __name__)
client = OpenAIAssistantClient()

# Decorador para proteção de rotas que exigem login (qualquer tipo de usuário)
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Obter ID da sessão do cookie
        session_id = request.cookies.get('session_id')
        
        # Validar a sessão
        session_data = db.validate_session(session_id)
        
        if not session_data:
            # Redirecionar para a página de login se não estiver autenticado
            return redirect(url_for('main.login_page'))
        
        # Adicionar dados do usuário ao contexto
        g.user = session_data
        
        # Se estiver autenticado, prosseguir
        return f(*args, **kwargs)
    
    return decorated_function

# Decorador para proteção de rotas administrativas
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Obter ID da sessão do cookie
        session_id = request.cookies.get('session_id')
        
        # Validar a sessão
        session_data = db.validate_session(session_id)
        
        if not session_data or (not session_data.get('is_admin') and session_data.get('tipo') != 'administrador'):
            # Redirecionar para a página de login se não estiver autenticado
            return redirect(url_for('main.login_page'))
        
        # Adicionar dados do usuário ao contexto
        g.user = session_data
        
        # Se estiver autenticado como admin, prosseguir
        return f(*args, **kwargs)
    
    return decorated_function

# Decorador para autenticação via API key
def api_key_required(permissions='read'):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            # Tentar obter API key do header Authorization
            auth_header = request.headers.get('Authorization')
            api_key = None
            
            if auth_header and auth_header.startswith('Bearer '):
                api_key = auth_header[7:]  # Remove 'Bearer '
            elif auth_header and auth_header.startswith('ApiKey '):
                api_key = auth_header[7:]  # Remove 'ApiKey '
            else:
                # Tentar obter do parâmetro de query como fallback
                api_key = request.args.get('api_key')
            
            if not api_key:
                return jsonify({
                    'error': 'API key obrigatória',
                    'message': 'Forneça a API key no header Authorization: "Bearer sua_api_key" ou como parâmetro ?api_key='
                }), 401
            
            # Validar API key
            key_data = db.validate_api_key(api_key)
            
            if not key_data:
                return jsonify({
                    'error': 'API key inválida ou expirada'
                }), 401
            
            # Verificar permissões se necessário
            if permissions != 'read' and key_data.get('permissions') != permissions and key_data.get('permissions') != 'admin':
                return jsonify({
                    'error': 'Permissões insuficientes para esta operação'
                }), 403
            
            # Adicionar dados da API key ao contexto
            g.api_key_data = key_data
            
            return f(*args, **kwargs)
        
        return decorated_function
    return decorator

# Health check endpoint para Docker
@bp.route('/health')
def health_check():
    """Endpoint de health check para verificar se a aplicação está funcionando"""
    try:
        # Verificar conexão com o banco de dados
        conn = db.get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT 1")
        conn.close()
        
        return jsonify({
            'status': 'healthy',
            'timestamp': datetime.now().isoformat(),
            'service': 'chatbot-assistant'
        }), 200
    except Exception as e:
        return jsonify({
            'status': 'unhealthy',
            'error': str(e),
            'timestamp': datetime.now().isoformat(),
            'service': 'chatbot-assistant'
        }), 503

# Rota pública para listar assistants sem autenticação
@bp.route('/public/assistants', methods=['GET'])
def public_assistants():
    """Lista pública de assistentes sem autenticação - apenas para acesso direto aos dados"""
    try:
        # Verificar parâmetros da requisição
        search_term = request.args.get('search', '')
        
        # Validar e corrigir limit e offset
        try:
            limit = int(request.args.get('limit', 1000))  # Aumentado de 100 para 1000
            if limit <= 0:
                limit = 1000
        except (ValueError, TypeError):
            limit = 1000
            
        try:
            offset = int(request.args.get('offset', 0))
            if offset < 0:
                offset = 0
        except (ValueError, TypeError):
            offset = 0
        
        # Buscar assistentes do banco local
        assistants = db.get_assistants(
            search_term=search_term if search_term else None,
            limit=limit,
            offset=offset
        )
        
        # Contar total para paginação
        total_assistants = db.count_assistants(
            search_term=search_term if search_term else None
        )
        
        # Verificar se há mais páginas
        has_more = (offset + limit) < total_assistants
        
        return jsonify({
            'assistants': assistants,
            'total': total_assistants,
            'has_more': has_more,
            'source': 'public_local'
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Rota para listar assistants com autenticação via API key
@bp.route('/api/v1/assistants', methods=['GET'])
@api_key_required('read')
def api_list_assistants():
    """Lista de assistentes para acesso via API key"""
    try:
        # Verificar parâmetros da requisição
        search_term = request.args.get('search', '')
        
        # Validar e corrigir limit e offset
        try:
            limit = int(request.args.get('limit', 1000))  # Aumentado de 100 para 1000
            if limit <= 0:
                limit = 1000
        except (ValueError, TypeError):
            limit = 1000
            
        try:
            offset = int(request.args.get('offset', 0))
            if offset < 0:
                offset = 0
        except (ValueError, TypeError):
            offset = 0
        
        # Buscar assistentes do banco local
        assistants = db.get_assistants(
            search_term=search_term if search_term else None,
            limit=limit,
            offset=offset
        )
        
        # Contar total para paginação
        total_assistants = db.count_assistants(
            search_term=search_term if search_term else None
        )
        
        # Verificar se há mais páginas
        has_more = (offset + limit) < total_assistants
        
        # Informações da API key usada (para logs/auditoria)
        api_key_info = {
            'key_name': g.api_key_data.get('key_name'),
            'permissions': g.api_key_data.get('permissions')
        }
        
        return jsonify({
            'assistants': assistants,
            'total': total_assistants,
            'has_more': has_more,
            'source': 'api_key_auth',
            'api_key_info': api_key_info
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Middleware para forçar HTTPS em ambientes de produção
@bp.before_request
def force_https():
    # Verificar se estamos atrás de um proxy HTTPS
    proto = request.headers.get('X-Forwarded-Proto', '')
    
    # Log da informação do protocolo para debug
    current_app.logger.info(f"X-Forwarded-Proto: {proto}")
    current_app.logger.info(f"Request path: {request.path}")
    current_app.logger.info(f"Request host: {request.host}")
    
    # Se estamos usando Cloudflare, configurar variáveis de contexto
    if 'X-Forwarded-Proto' in request.headers:
        g.is_https = (proto == 'https')
        g.scheme = proto
    else:
        g.is_https = request.is_secure
        g.scheme = 'https' if request.is_secure else 'http'
    
    # Definir a URL base para uso em templates e JavaScript
    g.base_url = f"{g.scheme}://{request.host}"
    
    current_app.logger.info(f"Base URL: {g.base_url}")

@bp.route('/home')
def home():
    """Página inicial que redireciona baseado no status de autenticação"""
    session_id = request.cookies.get('session_id')
    session_data = db.validate_session(session_id)
    
    if session_data:
        # Se estiver logado, redirecionar para a página apropriada
        if session_data.get('tipo') == 'administrador':
            return redirect(url_for('main.admin_assistants'))
        else:
            return redirect(url_for('main.dashboard'))
    else:
        # Se não estiver logado, redirecionar para login
        return redirect(url_for('main.login_page'))

@bp.route('/')
@login_required
def index():
    # Redirecionar baseado no tipo de usuário
    if g.user.get('tipo') == 'administrador':
        return redirect(url_for('main.admin_assistants'))
    else:
        return redirect(url_for('main.dashboard'))

@bp.route('/chat')
@login_required
def chat():
    """Página de chat com assistentes"""
    return render_template('index.html')

@bp.route('/login')
def login_page():
    # Verificar se já está logado
    session_id = request.cookies.get('session_id')
    session_data = db.validate_session(session_id)
    
    if session_data:
        # Se já estiver logado, redirecionar baseado no tipo
        if session_data.get('tipo') == 'administrador':
            return redirect(url_for('main.admin_assistants'))
        else:
            return redirect(url_for('main.dashboard'))
    
    return render_template('login.html')

@bp.route('/auth/login', methods=['POST'])
def login():
    try:
        data = request.json
        username = data.get('username')
        password = data.get('password')
        
        if not username or not password:
            return jsonify({'error': 'Nome de usuário e senha são obrigatórios'}), 400
        
        # Autenticar usuário
        user = db.authenticate_user(username, password)
        
        if not user:
            return jsonify({'error': 'Credenciais inválidas'}), 401
        
        # Criar sessão
        ip_address = request.remote_addr
        user_agent = request.headers.get('User-Agent')
        session_id = db.create_session(user['id'], ip_address, user_agent)
        
        # Determinar redirecionamento baseado no tipo de usuário
        redirect_url = '/admin/assistants' if user.get('tipo') == 'administrador' else '/dashboard'
        
        # Criar resposta com cookie
        response = jsonify({
            'message': 'Login realizado com sucesso',
            'user_type': user.get('tipo', 'usuario'),
            'redirect_url': redirect_url
        })
        response.set_cookie('session_id', session_id, httponly=True, max_age=86400)  # 24 horas
        
        return response
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/auth/logout')
def logout():
    # Obter ID da sessão do cookie
    session_id = request.cookies.get('session_id')
    
    if session_id:
        # Remover sessão do banco de dados
        db.delete_session(session_id)
    
    # Redirecionar para a página de login
    response = make_response(redirect(url_for('main.login_page')))
    response.delete_cookie('session_id')
    
    return response

@bp.route('/admin/assistants')
@admin_required
def admin_assistants():
    return render_template('admin/assistants.html')

@bp.route('/admin/settings')
@admin_required
def admin_settings():
    # Obter configurações atuais do banco de dados
    current_settings = db.get_system_settings()
    return render_template('admin/settings.html', settings=current_settings)

@bp.route('/admin/api/settings', methods=['GET'])
@admin_required
def api_get_settings():
    try:
        settings = db.get_system_settings()
        return jsonify(settings)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/settings', methods=['POST'])
@admin_required
def api_update_settings():
    try:
        data = request.json
        
        # Validar cores (formato hexadecimal)
        color_fields = ['primary_color', 'primary_dark', 'primary_light', 'secondary_color', 'accent_color']
        for field in color_fields:
            if field in data:
                color = data[field]
                if not color.startswith('#') or len(color) != 7:
                    return jsonify({'error': f'Cor {field} deve estar no formato #RRGGBB'}), 400
        
        # Atualizar configurações no banco
        db.update_system_settings(data)
        
        return jsonify({'message': 'Configurações atualizadas com sucesso'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/chat', methods=['POST'])
@login_required
def chat_with_assistant():
    try:
        data = request.json
        assistant_id = data.get('assistant_id')
        message = data.get('message')
        user_id = data.get('user_id', g.user['user_id'])  # Usar ID do usuário logado

        if not assistant_id or not message:
            return jsonify({
                'error': 'assistant_id e message são campos obrigatórios'
            }), 400

        # Adicionar logs para depuração
        current_app.logger.info(f"Chat request: assistant_id={assistant_id}, message={message[:50]}...")
        
        response = client.chat(assistant_id, message, user_id)
        
        # Adicionar logs para depuração da resposta
        current_app.logger.info(f"Chat response received, length: {len(response) if response else 0}")
        
        return jsonify({'response': response})

    except Exception as e:
        # Log detalhado do erro
        import traceback
        error_details = traceback.format_exc()
        current_app.logger.error(f"Erro no chat: {str(e)}")
        current_app.logger.error(f"Detalhes: {error_details}")
        
        return jsonify({'error': str(e)}), 500

@bp.route('/reset-thread', methods=['POST'])
@login_required
def reset_thread():
    try:
        data = request.json
        user_id = g.user['user_id']  # Usar ID do usuário logado
        client.reset_thread(user_id)
        return jsonify({'message': 'Thread resetada com sucesso'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500 

@bp.route('/list-assistants', methods=['GET'])
@login_required
def public_list_assistants():
    """Lista de assistentes para usuários autenticados"""
    try:
        # Verificar parâmetros da requisição
        search_term = request.args.get('search', '')
        
        # Validar e corrigir limit e offset
        try:
            limit = int(request.args.get('limit', 1000))  # Aumentado de 100 para 1000
            if limit <= 0:
                limit = 1000
        except (ValueError, TypeError):
            limit = 1000
            
        try:
            offset = int(request.args.get('offset', 0))
            if offset < 0:
                offset = 0
        except (ValueError, TypeError):
            offset = 0
        
        # Administradores têm acesso a todos os assistentes (não precisam de carteira)
        if g.user.get('tipo') == 'administrador' or g.user.get('is_admin'):
            # Buscar todos os assistentes do banco local
            all_assistants = db.get_assistants(
                search_term=search_term if search_term else None,
                limit=limit,
                offset=offset
            )
            
            # Contar total para paginação
            total_assistants = db.count_assistants(
                search_term=search_term if search_term else None
            )
            
            # Verificar se há mais páginas
            has_more = (offset + limit) < total_assistants
            
            return jsonify({
                'assistants': all_assistants,
                'total': total_assistants,
                'has_more': has_more,
                'source': 'admin_all'
            })
        
        # Para usuários comuns, buscar assistentes baseado na carteira
        user_assistants = db.get_user_assistants(g.user['user_id'])
        
        # Aplicar filtro de pesquisa se fornecido
        if search_term:
            filtered_assistants = []
            search_lower = search_term.lower()
            for assistant in user_assistants:
                if (search_lower in (assistant.get('name') or '').lower() or
                    search_lower in (assistant.get('description') or '').lower() or
                    search_lower in (assistant.get('id') or '').lower()):
                    filtered_assistants.append(assistant)
            user_assistants = filtered_assistants
        
        # Aplicar paginação
        total_assistants = len(user_assistants)
        paginated_assistants = user_assistants[offset:offset + limit]
        
        # Verificar se há mais páginas
        has_more = (offset + limit) < total_assistants
        
        return jsonify({
            'assistants': paginated_assistants,
            'total': total_assistants,
            'has_more': has_more,
            'source': 'wallet'
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/create-assistant', methods=['POST'])
@admin_required
def create_assistant():
    try:
        data = request.json
        model = data.get('model')
        name = data.get('name', None)
        description = data.get('description', None)
        instructions = data.get('instructions', None)
        tools = data.get('tools', [])

        # Faça a solicitação para a API da OpenAI
        response = requests.post(
            f"{client.base_url}/assistants",
            headers=client.headers,
            json={
                "model": model,
                "name": name,
                "description": description,
                "instructions": instructions,
                "tools": tools
            }
        )

        if response.status_code == 200:
            assistant_data = response.json()
            # Armazenar o assistente no banco de dados local
            db.store_assistant(assistant_data)
            return jsonify(assistant_data)
        else:
            return jsonify({'error': response.json().get('error', 'Erro ao criar assistente')}), response.status_code

    except Exception as e:
        return jsonify({'error': str(e)}), 500 

@bp.route('/admin/get-models', methods=['GET'])
@admin_required
def get_models():
    try:
        response = requests.get(
            f"{client.base_url}/models",
            headers=client.headers
        )

        if response.status_code == 200:
            models = response.json().get('data', [])
            return jsonify({'models': models})
        else:
            return jsonify({'error': response.json().get('error', 'Erro ao obter modelos')}), response.status_code

    except Exception as e:
        return jsonify({'error': str(e)}), 500 

@bp.route('/admin/list-assistants', methods=['GET'])
@admin_required
def list_assistants():
    try:
        # Verificar parâmetros da requisição
        search_term = request.args.get('search', '')
        
        # Validar e corrigir limit e offset
        try:
            limit = int(request.args.get('limit', 1000))  # Aumentado de 100 para 1000
            if limit <= 0:
                limit = 1000
        except (ValueError, TypeError):
            limit = 1000
            
        try:
            offset = int(request.args.get('offset', 0))
            if offset < 0:
                offset = 0
        except (ValueError, TypeError):
            offset = 0
            
        fetch_remote = request.args.get('fetch_remote', 'false').lower() == 'true'
        
        # Desativar sincronização automática completamente
        auto_sync = False
        
        # Primeiro, buscar do banco de dados local
        local_assistants = db.get_assistants(
            search_term=search_term if search_term else None, 
            limit=limit, 
            offset=offset
        )
        
        # Contar total de assistentes para paginação
        total_assistants = db.count_assistants(
            search_term=search_term if search_term else None
        )
        
        # Só buscar remotamente se explicitamente solicitado
        if fetch_remote:
            # Obter o after_id da API para paginação
            after = request.args.get('after', None)
            
            params = {
                'order': request.args.get('order', 'desc'),
                'limit': limit
            }
            
            # Adicionar o parâmetro 'after' se fornecido
            if after:
                params['after'] = after
            
            # Remover parâmetros nulos
            params = {k: v for k, v in params.items() if v is not None}
            
            # Buscar da API
            response = requests.get(
                f"{client.base_url}/assistants",
                headers=client.headers,
                params=params
            )
            
            if response.status_code != 200:
                return jsonify({'error': response.json().get('error', 'Erro ao listar assistentes')}), response.status_code
            
            result = response.json()
            
            # Armazenar os assistentes no banco de dados local
            for assistant in result.get('data', []):
                db.store_assistant(assistant)
            
            # Se tem um termo de pesquisa, filtrar os resultados da API
            api_assistants = result.get('data', [])
            if search_term:
                api_assistants = [a for a in api_assistants if 
                                 search_term.lower() in (a.get('name') or '').lower() or
                                 search_term.lower() in (a.get('description') or '').lower() or
                                 search_term.lower() in (a.get('id') or '').lower()]
            
            return jsonify({
                'assistants': api_assistants,
                'total': len(api_assistants),  # Para API, usar o tamanho da lista filtrada
                'has_more': result.get('has_more', False),
                'first_id': result.get('first_id', None),
                'last_id': result.get('last_id', None),
                'source': 'api'
            })
        
        # Retornar apenas os assistentes do banco de dados local
        return jsonify({
            'assistants': local_assistants,
            'total': total_assistants,
            'has_more': len(local_assistants) >= limit,
            'source': 'local'
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/sync-assistants', methods=['POST'])
@admin_required
def sync_assistants():
    """Sincroniza assistentes locais com a API da OpenAI"""
    try:
        # Obter informações da última sincronização
        sync_info = db.get_last_sync_info()
        cursor_after = sync_info.get('cursor_after')
        
        # Contar quantos assistentes foram sincronizados
        count = 0
        total_count = 0
        has_more = True
        
        # Buscar todos os assistentes usando paginação
        while has_more:
            params = {
                'order': 'desc',
                'limit': 100  # Máximo permitido pela API
            }
            
            if cursor_after:
                params['after'] = cursor_after
            
            response = requests.get(
                f"{client.base_url}/assistants",
                headers=client.headers,
                params=params
            )
            
            if response.status_code != 200:
                return jsonify({
                    'error': response.json().get('error', 'Erro ao sincronizar assistentes')
                }), response.status_code
            
            result = response.json()
            assistants = result.get('data', [])
            
            # Se não tiver assistentes, interrompe o loop
            if not assistants:
                has_more = False
                continue
            
            # Armazenar cada assistente no banco de dados
            for assistant in assistants:
                db.store_assistant(assistant)
                count += 1
            
            total_count += len(assistants)
            has_more = result.get('has_more', False)
            
            # Atualizar o cursor para a próxima página
            if has_more and result.get('last_id'):
                cursor_after = result.get('last_id')
            else:
                has_more = False
        
        # Atualizar informações de sincronização
        db.update_sync_info(cursor_after)
        
        return jsonify({
            'message': f'Sincronização concluída. {count} assistentes atualizados.',
            'count': count,
            'total_count': total_count
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/modify-assistant/<assistant_id>', methods=['POST'])
@admin_required
def modify_assistant(assistant_id):
    try:
        data = request.json
        model = data.get('model', None)
        name = data.get('name', None)
        description = data.get('description', None)
        instructions = data.get('instructions', None)

        # Faça a solicitação para a API da OpenAI
        response = requests.post(
            f"{client.base_url}/assistants/{assistant_id}",
            headers=client.headers,
            json={
                "model": model,
                "name": name,
                "description": description,
                "instructions": instructions
            }
        )

        if response.status_code == 200:
            # Atualizar o assistente no banco de dados local
            assistant_data = response.json()
            db.store_assistant(assistant_data)
            return jsonify({'message': 'Assistente atualizado com sucesso!', 'assistant': assistant_data})
        else:
            return jsonify({'error': response.json().get('error', 'Erro ao atualizar assistente')}), response.status_code

    except Exception as e:
        return jsonify({'error': str(e)}), 500 

@bp.route('/chat/send-message', methods=['POST'])
@login_required
def send_message():
    try:
        data = request.json
        message = data.get('message')
        assistant_id = data.get('assistant_id')  # Receber o ID do assistente selecionado
        
        if not assistant_id:
            return jsonify({'error': 'Nenhum assistente selecionado'}), 400

        # Use o assistant_id na chamada da API
        # ... resto do código ...

    except Exception as e:
        return jsonify({'error': str(e)}), 500 

@bp.route('/admin/delete-assistant/<assistant_id>', methods=['DELETE'])
@admin_required
def delete_assistant(assistant_id):
    try:
        response = requests.delete(
            f"{client.base_url}/assistants/{assistant_id}",
            headers=client.headers
        )

        if response.status_code == 200:
            # Marcar o assistente como excluído no banco de dados local
            db.mark_assistant_deleted(assistant_id)
            return jsonify({'message': 'Assistente deletado com sucesso!'})
        else:
            return jsonify({'error': response.json().get('error', 'Erro ao deletar assistente')}), response.status_code

    except Exception as e:
        return jsonify({'error': str(e)}), 500 

@bp.route('/admin/check-vector-store/<assistant_id>')
@admin_required
def check_vector_store(assistant_id):
    try:
        # Verificar o assistente
        response = requests.get(
            f"{client.base_url}/assistants/{assistant_id}",
            headers=client.headers
        )
        
        if not response.ok:
            return jsonify({'error': 'Erro ao verificar assistente'}), response.status_code

        assistant_data = response.json()
        tool_resources = assistant_data.get('tool_resources', {})
        file_search = tool_resources.get('file_search', {})
        vector_store_ids = file_search.get('vector_store_ids', [])

        if not vector_store_ids:
            return jsonify({'hasVectorStore': False})

        # Verificar o vector store
        vector_store_id = vector_store_ids[0]
        vector_response = requests.get(
            f"{client.base_url}/vector_stores/{vector_store_id}",
            headers=client.headers
        )

        if not vector_response.ok:
            return jsonify({'error': 'Erro ao verificar vector store'}), vector_response.status_code

        vector_data = vector_response.json()
        return jsonify({
            'hasVectorStore': True,
            'vectorStoreId': vector_store_id,
            'vectorStoreName': vector_data.get('name')
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/create-vector-store', methods=['POST'])
@admin_required
def create_vector_store():
    try:
        data = request.json
        name = data.get('name')
        assistant_id = data.get('assistantId')

        # Criar vector store
        response = requests.post(
            f"{client.base_url}/vector_stores",
            headers=client.headers,
            json={'name': name}
        )

        if not response.ok:
            return jsonify({'error': 'Erro ao criar vector store'}), response.status_code

        vector_store_data = response.json()
        
        # Atualizar o assistente com o novo vector store
        update_response = requests.post(
            f"{client.base_url}/assistants/{assistant_id}",
            headers=client.headers,
            json={
                'tool_resources': {
                    'file_search': {
                        'vector_store_ids': [vector_store_data['id']]
                    }
                }
            }
        )

        if not update_response.ok:
            return jsonify({'error': 'Erro ao atualizar assistente'}), update_response.status_code

        return jsonify({
            'vectorStoreId': vector_store_data['id'],
            'message': 'Vector store criado com sucesso'
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/chat/async', methods=['POST'])
@login_required
def chat_with_assistant_async():
    """Endpoint para processamento assíncrono ou síncrono de mensagens do chat, dependendo do ambiente."""
    try:
        data = request.json
        assistant_id = data.get('assistant_id')
        message = data.get('message')
        thread_id = data.get('thread_id')
        user_id = data.get('user_id', g.user['user_id'])  # Usar ID do usuário logado

        if not assistant_id or not message:
            return jsonify({
                'error': 'assistant_id e message são campos obrigatórios'
            }), 400

        # Verificar se deve usar processamento assíncrono ou síncrono
        use_async = current_app.config.get('USE_ASYNC', False)

        if use_async:
            # Modo assíncrono (Docker com Redis/Celery)
            try:
                # Importar a tarefa do Celery
                from app.tasks import process_chat_message
                
                # Submeter tarefa para processamento assíncrono
                task = process_chat_message.delay(assistant_id, message, thread_id)
                
                # Retornar imediatamente o ID da tarefa
                return jsonify({
                    'task_id': task.id,
                    'status': 'processing'
                })
            except Exception as e:
                # Fallback para modo síncrono se Celery falhar
                print(f"Erro ao usar Celery, caindo para modo síncrono: {str(e)}")
                use_async = False

        # Modo síncrono (desenvolvimento local sem Redis)
        if not use_async:
            # Usar o cliente OpenAI diretamente
            response = client.chat(assistant_id, message, user_id)
            return jsonify({
                'ready': True,
                'status': 'concluído',
                'result': {
                    'response': response,
                    'thread_id': client.active_threads.get(user_id)
                }
            })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/chat/status/<task_id>', methods=['GET'])
def check_chat_status(task_id):
    """Endpoint para verificar o status de uma tarefa de chat."""
    try:
        from app.tasks import celery_app
        
        # Verificar o status da tarefa
        task = celery_app.AsyncResult(task_id)
        
        if task.state == 'PENDING':
            response = {
                'status': 'em andamento',
                'ready': False
            }
        elif task.state == 'FAILURE':
            response = {
                'status': 'falhou',
                'error': str(task.info),
                'ready': True
            }
        elif task.state == 'SUCCESS':
            response = {
                'status': 'concluído',
                'result': task.result,
                'ready': True
            }
        else:
            response = {
                'status': task.state,
                'ready': False
            }
        
        return jsonify(response)
    except Exception as e:
        return jsonify({'error': str(e)}), 500 

@bp.route('/admin/users')
@admin_required
def admin_users():
    return render_template('admin/users.html')

@bp.route('/admin/api/users')
@admin_required
def api_list_users():
    try:
        users = db.get_all_users()
        return jsonify({'users': users})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/users/<int:user_id>')
@admin_required
def api_get_user(user_id):
    try:
        user = db.get_user(user_id)
        if not user:
            return jsonify({'error': 'Usuário não encontrado'}), 404
        return jsonify(user)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/users', methods=['POST'])
@admin_required
def api_create_user():
    """Cria um novo usuário"""
    try:
        data = request.json
        username = data.get('username')
        password = data.get('password')
        email = data.get('email')
        full_name = data.get('full_name')
        is_admin = data.get('is_admin', 0)
        wallet_id = data.get('wallet_id')
        
        if not username or not password:
            return jsonify({'error': 'Username e senha são obrigatórios'}), 400
        
        success, result = db.create_user(username, password, email, full_name, is_admin)
        
        if success:
            user_id = result
            
            # Atribuir carteira apenas se não for administrador
            if wallet_id and not is_admin:
                wallet_success, wallet_message = db.assign_wallet_to_user(user_id, wallet_id)
                if not wallet_success:
                    # Se falhar ao atribuir carteira, ainda retorna sucesso mas com aviso
                    return jsonify({
                        'message': 'Usuário criado com sucesso, mas houve erro ao atribuir carteira',
                        'warning': wallet_message
                    }), 201
            
            return jsonify({'message': 'Usuário criado com sucesso'}), 201
        else:
            return jsonify({'error': result}), 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/users/<int:user_id>', methods=['PUT'])
@admin_required
def api_update_user(user_id):
    """Atualiza um usuário"""
    try:
        data = request.json
        wallet_id = data.get('wallet_id')
        is_admin = data.get('is_admin', 0)
        
        # Remover wallet_id dos dados do usuário para não interferir na atualização
        user_data = {k: v for k, v in data.items() if k != 'wallet_id'}
        
        success, message = db.update_user(user_id, user_data)
        
        if success:
            # Atualizar carteira apenas se não for administrador
            if 'wallet_id' in data and not is_admin:
                wallet_success, wallet_message = db.assign_wallet_to_user(user_id, wallet_id)
                if not wallet_success:
                    return jsonify({
                        'message': 'Usuário atualizado, mas houve erro ao atribuir carteira',
                        'warning': wallet_message
                    })
            elif is_admin:
                # Se virou administrador, remover qualquer carteira existente
                db.assign_wallet_to_user(user_id, None)
            
            return jsonify({'message': message})
        else:
            return jsonify({'error': message}), 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/users/<int:user_id>', methods=['DELETE'])
@admin_required
def api_delete_user(user_id):
    try:
        success, message = db.delete_user(user_id)
        
        if success:
            return jsonify({'message': message})
        else:
            return jsonify({'error': message}), 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500 

@bp.route('/admin/upload-files', methods=['POST'])
@admin_required
def upload_files():
    try:
        files = request.files.getlist('files')
        vector_store_id = request.form.get('vectorStoreId')
        
        if not files:
            return jsonify({'error': 'Nenhum arquivo enviado'}), 400

        uploaded_file_ids = []

        # Upload dos arquivos para a API de Files
        for file in files:
            files_response = requests.post(
                f"{client.base_url}/files",
                headers={
                    "Authorization": client.headers['Authorization'],
                },
                files={
                    'file': (file.filename, file.stream, file.content_type),
                    'purpose': (None, 'assistants')
                }
            )

            if not files_response.ok:
                error_data = files_response.json()
                # Verificar se o erro é relacionado a arquivo não legível
                if 'not readable' in str(error_data.get('error', '')).lower():
                    return jsonify({
                        'error': 'O arquivo não está legível, você deve tornar o arquivo um documento com texto localizável para poder realizar o upload'
                    }), 400
                return jsonify({'error': f'Erro ao fazer upload do arquivo {file.filename}'}), files_response.status_code

            file_data = files_response.json()
            uploaded_file_ids.append(file_data['id'])
            file.stream.seek(0)

        # Agora, criar um batch com os arquivos no vector store
        batch_response = requests.post(
            f"{client.base_url}/vector_stores/{vector_store_id}/file_batches",
            headers=client.headers,
            json={
                'file_ids': uploaded_file_ids
            }
        )

        if not batch_response.ok:
            return jsonify({'error': 'Erro ao adicionar arquivos ao vector store'}), batch_response.status_code

        batch_data = batch_response.json()

        # Aguardar o processamento dos arquivos (com timeout de 1 hora)
        max_wait_time = 3600  # 60 minutos (1 hora) para arquivos grandes
        start_time = time.time()

        while time.time() - start_time < max_wait_time:
            status_response = requests.get(
                f"{client.base_url}/vector_stores/{vector_store_id}/file_batches/{batch_data['id']}",
                headers=client.headers
            )

            if not status_response.ok:
                return jsonify({'error': 'Erro ao verificar status do upload'}), status_response.status_code

            status_data = status_response.json()
            if status_data['status'] in ['completed', 'failed']:
                break

            time.sleep(1)  # Aguarda 1 segundo antes de verificar novamente

        # Verificar se deu timeout
        if status_data['status'] not in ['completed', 'failed']:
            return jsonify({'error': f'Timeout após {max_wait_time} segundos. O processamento continua em background.'}), 408

        return jsonify({
            'status': status_data['status'],
            'file_counts': status_data['file_counts'],
            'message': 'Arquivos enviados com sucesso'
        })

    except Exception as e:
        if 'not readable' in str(e).lower():
            return jsonify({
                'error': 'O arquivo não está legível, você deve tornar o arquivo um documento com texto localizável para poder realizar o upload'
            }), 400
        return jsonify({'error': str(e)}), 500 

# Rota para gerar documentos a partir dos dados do fluxo guiado
@bp.route('/gerar-documento', methods=['POST'])
def gerar_documento():
    # Obter dados enviados pelo cliente
    data = request.json
    tipo_documento = data.get('tipo')
    dados = data.get('dados')
    assistant_id = data.get('assistant_id')
    user_id = data.get('user_id')
    
    # Validar dados
    if not tipo_documento or not dados:
        return jsonify({'error': 'Dados incompletos'}), 400
    
    try:
        # Criar diretório para documentos se não existir
        docs_dir = os.path.join(current_app.root_path, 'static', 'documentos')
        if not os.path.exists(docs_dir):
            os.makedirs(docs_dir)
        
        # Gerar nome de arquivo único
        documento_id = str(uuid.uuid4())
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        
        # Tentar gerar texto com IA se tiver assistant_id
        texto_gerado = None
        nome_assistant = None
        
        if assistant_id:
            try:
                # Buscar nome do assistant
                assistant = db.get_assistant_by_id(assistant_id)
                if assistant:
                    nome_assistant = assistant.get('name', 'Condomínio')
                
                # Preparar dados para IA
                ocorrencia = {
                    'external_assistant_id': assistant_id,
                    'morador': {
                        'nome': f"Morador(a) responsável",
                        'apartamento': 'Unidade não especificada',
                        'bloco': ''
                    },
                    'data': dados.get('data'),
                    'descricao': dados.get('descricao', ''),
                    'valor': float(dados.get('valor', 0)) if tipo_documento == 'multa' else None
                }
                
                # Importar e usar o serviço
                from app.services.openai_service import OpenAIService
                import asyncio
                
                service = OpenAIService()
                
                # Executar a função assíncrona
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                texto_gerado = loop.run_until_complete(
                    service.gerar_documento_com_assistant(ocorrencia, tipo_documento)
                )
                loop.close()
                
            except Exception as e:
                current_app.logger.error(f"Erro ao gerar com IA: {str(e)}")
                # Continuar com geração padrão
        
        # Determinar tipo de documento e gerar arquivo
        if tipo_documento == 'advertencia':
            documento_url = gerar_advertencia_aprimorada(
                documento_id, timestamp, dados, docs_dir, 
                texto_gerado, nome_assistant
            )
        elif tipo_documento == 'multa':
            documento_url = gerar_multa_aprimorada(
                documento_id, timestamp, dados, docs_dir,
                texto_gerado, nome_assistant
            )
        else:
            return jsonify({'error': 'Tipo de documento não suportado'}), 400
        
        # Retornar URL do documento gerado
        return jsonify({
            'success': True,
            'documento_url': documento_url
        })
        
    except Exception as e:
        current_app.logger.error(f"Erro ao gerar documento: {str(e)}")
        return jsonify({'error': str(e)}), 500

# Função para gerar documento de advertência como HTML
def gerar_advertencia(documento_id, timestamp, dados, docs_dir):
    try:
        # Definir nome do arquivo
        filename = f"advertencia_{timestamp}_{documento_id[:8]}.html"
        filepath = os.path.join(docs_dir, filename)
        
        # Formatação da data para exibição
        data_incidente = dados.get('data', 'Não informada')
        descricao = dados.get('descricao', 'Não informada')
        
        # Criar conteúdo HTML
        html_content = f"""
        <!DOCTYPE html>
        <html lang="pt-BR">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Advertência Disciplinar</title>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    margin: 40px;
                    line-height: 1.6;
                }}
                .header {{
                    text-align: center;
                    margin-bottom: 30px;
                }}
                .title {{
                    font-size: 24px;
                    font-weight: bold;
                    margin-bottom: 20px;
                    text-transform: uppercase;
                }}
                .section {{
                    margin-bottom: 20px;
                }}
                .section-title {{
                    font-weight: bold;
                    font-size: 16px;
                    margin-bottom: 10px;
                }}
                .footer {{
                    margin-top: 50px;
                    font-style: italic;
                    font-size: 12px;
                    color: #666;
                    text-align: center;
                }}
                .signature {{
                    margin-top: 80px;
                    border-top: 1px solid #000;
                    width: 200px;
                    text-align: center;
                    padding-top: 10px;
                }}
                @media print {{
                    body {{
                        margin: 0;
                        padding: 20px;
                    }}
                    .no-print {{
                        display: none;
                    }}
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="title">Advertência Disciplinar</div>
            </div>
            
            <div class="section">
                <div class="section-title">Data do Incidente:</div>
                <div>{data_incidente}</div>
            </div>
            
            <div class="section">
                <div class="section-title">Descrição do Incidente:</div>
                <div>{descricao}</div>
            </div>
            
            <div class="signature">
                Assinatura do Responsável
            </div>
            
            <div class="footer">
                Este documento foi gerado automaticamente pelo sistema em {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}.
            </div>
            
            <div class="no-print" style="margin-top: 30px; text-align: center;">
                <button onclick="window.print()">Imprimir Documento</button>
            </div>
        </body>
        </html>
        """
        
        # Escrever conteúdo HTML no arquivo
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # Retornar URL relativa para o documento
        return f"/static/documentos/{filename}"
        
    except Exception as e:
        current_app.logger.error(f"Erro ao gerar advertência: {str(e)}")
        raise
        
# Função para gerar documento de multa como HTML
def gerar_multa(documento_id, timestamp, dados, docs_dir):
    try:
        # Definir nome do arquivo
        filename = f"multa_{timestamp}_{documento_id[:8]}.html"
        filepath = os.path.join(docs_dir, filename)
        
        # Obter dados
        data_infracao = dados.get('data', 'Não informada')
        valor_multa = dados.get('valor', '0,00')
        descricao = dados.get('descricao', 'Não informada')
        
        # Criar conteúdo HTML
        html_content = f"""
        <!DOCTYPE html>
        <html lang="pt-BR">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Notificação de Multa</title>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    margin: 40px;
                    line-height: 1.6;
                }}
                .header {{
                    text-align: center;
                    margin-bottom: 30px;
                }}
                .title {{
                    font-size: 24px;
                    font-weight: bold;
                    margin-bottom: 20px;
                    text-transform: uppercase;
                }}
                .section {{
                    margin-bottom: 20px;
                }}
                .section-title {{
                    font-weight: bold;
                    font-size: 16px;
                    margin-bottom: 10px;
                }}
                .info-table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin-bottom: 20px;
                }}
                .info-table th, .info-table td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                }}
                .info-table th {{
                    background-color: #f2f2f2;
                    text-align: left;
                    width: 30%;
                }}
                .footer {{
                    margin-top: 50px;
                    font-style: italic;
                    font-size: 12px;
                    color: #666;
                    text-align: center;
                }}
                .signature {{
                    margin-top: 80px;
                    border-top: 1px solid #000;
                    width: 200px;
                    text-align: center;
                    padding-top: 10px;
                }}
                .highlight {{
                    font-weight: bold;
                    color: #d9534f;
                }}
                @media print {{
                    body {{
                        margin: 0;
                        padding: 20px;
                    }}
                    .no-print {{
                        display: none;
                    }}
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="title">Notificação de Multa</div>
            </div>
            
            <table class="info-table">
                <tr>
                    <th>Data da Infração:</th>
                    <td>{data_infracao}</td>
                </tr>
                <tr>
                    <th>Valor da Multa:</th>
                    <td class="highlight">R$ {valor_multa}</td>
                </tr>
            </table>
            
            <div class="section">
                <div class="section-title">Descrição da Infração:</div>
                <div>{descricao}</div>
            </div>
            
            <div class="signature">
                Assinatura do Responsável
            </div>
            
            <div class="footer">
                Este documento foi gerado automaticamente pelo sistema em {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}.
            </div>
            
            <div class="no-print" style="margin-top: 30px; text-align: center;">
                <button onclick="window.print()">Imprimir Documento</button>
            </div>
        </body>
        </html>
        """
        
        # Escrever conteúdo HTML no arquivo
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # Retornar URL relativa para o documento
        return f"/static/documentos/{filename}"
        
    except Exception as e:
        current_app.logger.error(f"Erro ao gerar multa: {str(e)}")
        raise 

@bp.route('/api/gerar-documento', methods=['POST'])
def api_gerar_documento():
    """
    Endpoint API para ser chamado pelo assistente OpenAI como ferramenta
    """
    try:
        # Obter dados enviados pelo cliente
        data = request.json
        
        # Mostrar dados recebidos no log para depuração
        current_app.logger.info(f"Dados recebidos na API: {data}")
        
        tipo_documento = data.get('tipo')
        dados = {
            'data': data.get('data'),
            'descricao': data.get('descricao'),
            'valor': data.get('valor', '')  # Valor padrão vazio
        }
        
        # Validar dados
        if not tipo_documento or not dados.get('data') or not dados.get('descricao'):
            error_msg = {'error': 'Dados incompletos. Necessário tipo, data e descrição'}
            current_app.logger.error(f"Erro de validação: {error_msg}")
            return jsonify(error_msg), 400
        
        # Criar diretório para documentos se não existir
        docs_dir = os.path.join(current_app.root_path, 'static', 'documentos')
        if not os.path.exists(docs_dir):
            os.makedirs(docs_dir)
        
        # Gerar nome de arquivo único
        documento_id = str(uuid.uuid4())
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        
        # Determinar tipo de documento e gerar arquivo
        if tipo_documento.lower() == 'advertencia':
            documento_url = gerar_advertencia(documento_id, timestamp, dados, docs_dir)
        elif tipo_documento.lower() == 'multa':
            documento_url = gerar_multa(documento_id, timestamp, dados, docs_dir)
        else:
            error_msg = {'error': 'Tipo de documento não suportado'}
            current_app.logger.error(f"Tipo inválido: {tipo_documento}")
            return jsonify(error_msg), 400
        
        # URL completa para o documento - usar o mesmo protocolo e host da requisição
        # Isso garante que funcione tanto com http quanto https e com cloudflared
        request_host = request.headers.get('Host') or request.host
        request_protocol = 'https' if request.is_secure or 'cloudflare' in request_host else 'http'
        base_url = f"{request_protocol}://{request_host}"
        documento_url_completa = f"{base_url}{documento_url}"
        
        # Registrar a URL no log
        current_app.logger.info(f"URL do documento: {documento_url_completa}")
        
        # Criar resposta
        resposta = {
            'success': True,
            'mensagem': f'Documento de {tipo_documento} gerado com sucesso!',
            'documento_url': documento_url_completa,
            'tipo': tipo_documento
        }
        
        # Registrar resposta no log
        current_app.logger.info(f"Resposta da API: {resposta}")
        
        # Retornar resposta JSON
        return jsonify(resposta)
        
    except Exception as e:
        error_msg = {'error': str(e)}
        current_app.logger.error(f"Erro ao gerar documento: {str(e)}")
        return jsonify(error_msg), 500 

@bp.route('/create-thread-and-run', methods=['POST', 'OPTIONS'])
def create_thread_and_run():
    """
    Endpoint para criar uma thread e executar um assistente.
    Suporta requisições CORS com método OPTIONS.
    """
    # Tratar solicitações preflight (OPTIONS)
    if request.method == 'OPTIONS':
        response = jsonify({})
        return response
    
    try:
        data = request.json
        assistant_id = data.get('assistant_id')
        instructions = data.get('instructions')
        message = data.get('message')
        user_id = data.get('user_id', 'default_user')
        
        if not assistant_id or not message:
            return jsonify({
                'error': 'assistant_id e message são campos obrigatórios'
            }), 400
            
        # Verificar se o assistente existe
        assistant = db.get_assistant_by_id(assistant_id)
        if not assistant:
            return jsonify({'error': 'Assistente não encontrado'}), 404
        
        # Instrução específica do usuário, se fornecida
        user_instructions = None
        if instructions:
            user_instructions = instructions
        
        # Criar thread no serviço de chat
        response = client.chat(
            assistant_id=assistant_id, 
            message=message, 
            user_id=user_id, 
            instructions=user_instructions
        )
        
        # Obter o ID da thread atual
        thread_id = client.active_threads.get(user_id)
        
        return jsonify({
            'response': response,
            'thread_id': thread_id,
            'assistant_id': assistant_id
        })
        
    except Exception as e:
        # Log detalhado do erro
        import traceback
        error_details = traceback.format_exc()
        current_app.logger.error(f"Erro ao criar thread e executar: {str(e)}")
        current_app.logger.error(f"Detalhes: {error_details}")
        
        return jsonify({'error': str(e)}), 500 

@bp.route('/register')
def register_page():
    """Página de registro de novos usuários"""
    # Verificar se já está logado
    session_id = request.cookies.get('session_id')
    session_data = db.validate_session(session_id)
    
    if session_data:
        # Se já estiver logado, redirecionar baseado no tipo
        if session_data.get('tipo') == 'administrador':
            return redirect(url_for('main.admin_assistants'))
        else:
            return redirect(url_for('main.dashboard'))
    
    return render_template('register.html')

@bp.route('/auth/register', methods=['POST'])
def register():
    """Endpoint para registrar novos usuários"""
    try:
        data = request.json
        username = data.get('username')
        password = data.get('password')
        email = data.get('email')
        full_name = data.get('full_name')
        
        # Validações
        if not username or not password or not email or not full_name:
            return jsonify({'error': 'Todos os campos são obrigatórios'}), 400
        
        if len(password) < 6:
            return jsonify({'error': 'A senha deve ter pelo menos 6 caracteres'}), 400
        
        # Registrar usuário
        success, result = db.register_user(username, password, email, full_name)
        
        if not success:
            return jsonify({'error': result}), 400
        
        return jsonify({
            'message': 'Usuário registrado com sucesso!',
            'redirect_url': '/login'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/dashboard')
@login_required
def dashboard():
    """Dashboard do usuário"""
    # Gerar mensagem de saudação baseada no horário
    hour = datetime.now().hour
    
    if hour < 12:
        greeting_message = 'Bom dia! Como posso ajudá-lo hoje?'
    elif hour < 18:
        greeting_message = 'Boa tarde! Como posso ajudá-lo hoje?'
    else:
        greeting_message = 'Boa noite! Como posso ajudá-lo hoje?'
    
    return render_template('dashboard.html', user=g.user, greeting_message=greeting_message)

@bp.route('/profile')
@login_required
def profile():
    """Página de perfil do usuário"""
    # Obter nome da carteira do usuário
    wallet_name = db.get_user_wallet_name(g.user['user_id'])
    user_data = g.user.copy()
    user_data['wallet_name'] = wallet_name
    return render_template('profile.html', user=user_data)

@bp.route('/api/auth/check-session', methods=['GET'])
def check_session():
    """Verifica se há uma sessão ativa"""
    session_id = request.cookies.get('session_id')
    session_data = db.validate_session(session_id)
    
    if session_data:
        return jsonify({
            'authenticated': True,
            'user': {
                'id': session_data['user_id'],
                'username': session_data['username'],
                'email': session_data['email'],
                'full_name': session_data['full_name'],
                'tipo': session_data['tipo'],
                'carteira': session_data['carteira']
            }
        })
    
    return jsonify({'authenticated': False})

@bp.route('/api/user/profile', methods=['GET', 'PUT'])
@login_required
def user_profile():
    """Obter ou atualizar perfil do usuário"""
    if request.method == 'GET':
        user = db.get_user(g.user['user_id'])
        if user:
            return jsonify({
                'id': user['id'],
                'username': user['username'],
                'email': user['email'],
                'full_name': user['full_name'],
                'tipo': user['tipo'],
                'carteira': user['carteira'],
                'created_at': user['created_at'],
                'last_login': user['last_login'],
                'last_update': user['last_update']
            })
        return jsonify({'error': 'Usuário não encontrado'}), 404
    
    elif request.method == 'PUT':
        data = request.json
        allowed_fields = ['email', 'full_name', 'password']
        update_data = {k: v for k, v in data.items() if k in allowed_fields}
        
        success, message = db.update_user(g.user['user_id'], update_data)
        
        if success:
            return jsonify({'message': message})
        return jsonify({'error': message}), 400

@bp.route('/api/user/wallet', methods=['GET'])
@login_required
def user_wallet():
    """Obter informações da carteira do usuário"""
    user = db.get_user(g.user['user_id'])
    if user:
        return jsonify({
            'balance': user['carteira'],
            'currency': 'BRL'
        })
    return jsonify({'error': 'Usuário não encontrado'}), 404

# ============ ENDPOINTS PARA TOKENS DE INTEGRAÇÃO ============

@bp.route('/api/user/integration-token', methods=['GET'])
@login_required
def get_integration_token():
    """Obter token de integração do usuário atual"""
    try:
        token_data = db.get_user_integration_token(g.user['user_id'])
        
        if token_data:
            # Não retorna o token completo por segurança, apenas informações
            return jsonify({
                'has_token': True,
                'created_at': token_data['created_at'],
                'last_used': token_data['last_used'],
                'description': token_data['description'],
                'token_preview': token_data['token'][:8] + '...' + token_data['token'][-8:]
            })
        else:
            return jsonify({'has_token': False})
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/api/user/integration-token', methods=['POST'])
@login_required
def generate_user_integration_token():
    """Gerar novo token de integração para o usuário atual"""
    try:
        data = request.json
        description = data.get('description', 'Token gerado pelo usuário')
        
        token = db.generate_integration_token(g.user['user_id'], description=description)
        
        return jsonify({
            'message': 'Token de integração gerado com sucesso',
            'token': token,
            'instructions': 'Copie este token e configure-o no sistema que deseja integrar. Por segurança, o token não será mostrado novamente.'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/api/user/integration-token', methods=['DELETE'])
@login_required
def revoke_user_integration_token():
    """Revogar token de integração do usuário atual"""
    try:
        success = db.revoke_integration_token(g.user['user_id'])
        
        if success:
            return jsonify({'message': 'Token de integração revogado com sucesso'})
        else:
            return jsonify({'error': 'Erro ao revogar token'}), 400
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/auth/integration-login/<token>')
def integration_login(token):
    """Autenticação via token de integração"""
    print(f"=== INTEGRATION LOGIN DEBUG ===")
    print(f"Token recebido: {token}")
    print(f"IP: {request.remote_addr}")
    print(f"User-Agent: {request.headers.get('User-Agent')}")
    
    try:
        print("Tentando autenticar com token...")
        user = db.authenticate_by_integration_token(token)
        print(f"Resultado da autenticação: {user}")
        
        if user:
            print(f"Usuário autenticado: ID={user.get('id')}, Username={user.get('username')}")
            
            # Criar sessão normal
            print("Criando sessão...")
            session_id = db.create_session(
                user['id'], 
                request.remote_addr, 
                request.headers.get('User-Agent')
            )
            print(f"Sessão criada: {session_id}")
            
            # Redirecionar para dashboard
            print("Redirecionando para dashboard...")
            response = make_response(redirect(url_for('main.dashboard')))
            response.set_cookie('session_id', session_id, httponly=True, max_age=86400)
            
            return response
        else:
            print("Token inválido ou usuário não encontrado")
            # Token inválido, redirecionar para login normal
            return redirect(url_for('main.login_page', error='Token de integração inválido'))
            
    except Exception as e:
        print(f"ERRO: {str(e)}")
        print(f"Traceback: {e.__traceback__}")
        current_app.logger.error(f"Erro na autenticação por token de integração: {str(e)}")
        return redirect(url_for('main.login_page', error='Erro na autenticação'))

@bp.route('/admin/api/integration-tokens', methods=['GET'])
@admin_required
def admin_list_integration_tokens():
    """Listar todos os tokens de integração (apenas admin)"""
    try:
        tokens = db.get_all_integration_tokens()
        return jsonify({'tokens': tokens})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/users/<int:user_id>/integration-token', methods=['GET'])
@admin_required
def admin_get_user_integration_token(user_id):
    """Obter token de integração de um usuário específico (admin)"""
    try:
        token_data = db.get_user_integration_token(user_id)
        
        if token_data:
            return jsonify({
                'has_token': True,
                'created_at': token_data['created_at'],
                'last_used': token_data['last_used'],
                'description': token_data['description'],
                'token_preview': token_data['token'][:8] + '...' + token_data['token'][-8:]
            })
        else:
            return jsonify({'has_token': False})
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/users/<int:user_id>/integration-token', methods=['POST'])
@admin_required
def admin_generate_user_integration_token(user_id):
    """Gerar token de integração para um usuário específico (admin)"""
    try:
        data = request.json
        description = data.get('description', f'Token gerado pelo administrador para usuário ID {user_id}')
        
        token = db.generate_integration_token(user_id, description=description)
        
        return jsonify({
            'message': 'Token de integração gerado com sucesso',
            'token': token,
            'instructions': 'Copie este token e configure-o no sistema que deseja integrar. Por segurança, o token não será mostrado novamente.'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/users/<int:user_id>/integration-token', methods=['DELETE'])
@admin_required
def admin_revoke_user_integration_token(user_id):
    """Revogar token de integração de um usuário específico (admin)"""
    try:
        success = db.revoke_integration_token(user_id)
        
        if success:
            return jsonify({'message': 'Token de integração revogado com sucesso'})
        else:
            return jsonify({'error': 'Erro ao revogar token'}), 400
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Rotas para gerenciar carteiras
@bp.route('/admin/wallets')
@admin_required
def admin_wallets():
    """Página de administração de carteiras"""
    return render_template('admin/wallets.html')

@bp.route('/admin/api/wallets', methods=['GET'])
@admin_required
def api_list_wallets():
    """Lista todas as carteiras"""
    try:
        wallets = db.get_all_wallets()
        return jsonify({'wallets': wallets})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/wallets', methods=['POST'])
@admin_required
def api_create_wallet():
    """Cria uma nova carteira"""
    try:
        data = request.json
        name = data.get('name')
        description = data.get('description', '')
        
        if not name:
            return jsonify({'error': 'Nome da carteira é obrigatório'}), 400
        
        # Debug: verificar o conteúdo de g.user
        current_app.logger.info(f"g.user content: {g.user if hasattr(g, 'user') else 'g.user not set'}")
        
        # Obter user_id de forma segura
        user_id = None
        if hasattr(g, 'user') and g.user:
            user_id = g.user.get('user_id')
            current_app.logger.info(f"user_id found: {user_id}")
        
        success, result = db.create_wallet(name, description, user_id)
        
        if success:
            return jsonify({'message': 'Carteira criada com sucesso', 'wallet_id': result}), 201
        else:
            return jsonify({'error': result}), 400
        
    except Exception as e:
        current_app.logger.error(f"Erro ao criar carteira: {str(e)}")
        import traceback
        current_app.logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/wallets/<int:wallet_id>', methods=['GET'])
@admin_required
def api_get_wallet(wallet_id):
    """Obtém uma carteira específica"""
    try:
        wallet = db.get_wallet(wallet_id)
        if not wallet:
            return jsonify({'error': 'Carteira não encontrada'}), 404
        return jsonify(wallet)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/wallets/<int:wallet_id>', methods=['PUT'])
@admin_required
def api_update_wallet(wallet_id):
    """Atualiza uma carteira"""
    try:
        data = request.json
        name = data.get('name')
        description = data.get('description')
        
        success, message = db.update_wallet(wallet_id, name, description)
        
        if success:
            return jsonify({'message': message})
        else:
            return jsonify({'error': message}), 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/wallets/<int:wallet_id>', methods=['DELETE'])
@admin_required
def api_delete_wallet(wallet_id):
    """Remove uma carteira"""
    try:
        success, message = db.delete_wallet(wallet_id)
        
        if success:
            return jsonify({'message': message})
        else:
            return jsonify({'error': message}), 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/wallets/<int:wallet_id>/assistants', methods=['GET'])
@admin_required
def api_get_wallet_assistants(wallet_id):
    """Lista assistentes de uma carteira"""
    try:
        assistants = db.get_wallet_assistants(wallet_id)
        return jsonify({'assistants': assistants})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/wallets/<int:wallet_id>/assistants', methods=['POST'])
@admin_required
def api_add_assistant_to_wallet(wallet_id):
    """Adiciona um assistente à carteira"""
    try:
        data = request.json
        assistant_id = data.get('assistant_id')
        
        if not assistant_id:
            return jsonify({'error': 'ID do assistente é obrigatório'}), 400
        
        # Verificar se o assistente existe
        assistant = db.get_assistant_by_id(assistant_id)
        if not assistant:
            return jsonify({'error': 'Assistente não encontrado'}), 404
        
        # Obter user_id de forma segura
        user_id = None
        if hasattr(g, 'user') and g.user:
            user_id = g.user.get('user_id')
        
        success, message = db.add_assistant_to_wallet(wallet_id, assistant_id, user_id)
        
        if success:
            return jsonify({'message': message})
        else:
            return jsonify({'error': message}), 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/wallets/<int:wallet_id>/assistants/<assistant_id>', methods=['DELETE'])
@admin_required
def api_remove_assistant_from_wallet(wallet_id, assistant_id):
    """Remove um assistente da carteira"""
    try:
        success, message = db.remove_assistant_from_wallet(wallet_id, assistant_id)
        
        if success:
            return jsonify({'message': message})
        else:
            return jsonify({'error': message}), 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/users/<int:user_id>/wallet', methods=['PUT'])
@admin_required
def api_assign_wallet_to_user(user_id):
    try:
        data = request.get_json()
        wallet_id = data.get('wallet_id')
        
        if not wallet_id:
            return jsonify({'error': 'ID da carteira é obrigatório'}), 400
        
        success = db.assign_wallet_to_user(user_id, wallet_id)
        
        if success:
            return jsonify({'message': 'Carteira atribuída com sucesso'})
        else:
            return jsonify({'error': 'Erro ao atribuir carteira'}), 400
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ============ ENDPOINTS DE SEGURANÇA SQLITE ============

@bp.route('/admin/security')
@admin_required
def admin_security():
    """Página de administração de segurança do banco de dados"""
    return render_template('admin/security.html')

@bp.route('/admin/api/security/status', methods=['GET'])
@admin_required
def api_security_status():
    """Obter status de segurança do banco de dados"""
    try:
        from app import database as db
        
        # Verificar se o módulo de segurança está disponível
        security_enabled = hasattr(db, 'SECURITY_ENABLED') and db.SECURITY_ENABLED
        
        # Obter estatísticas do banco
        stats = db.get_database_stats()
        
        # Verificar integridade
        integrity_ok = db.check_database_integrity()
        
        # Verificar se existem backups
        backup_dir = os.path.join(os.path.dirname(db.DB_PATH), 'backups')
        backup_count = 0
        latest_backup = None
        
        if os.path.exists(backup_dir):
            backups = [f for f in os.listdir(backup_dir) if f.endswith('.db')]
            backup_count = len(backups)
            if backups:
                latest_backup = max(backups, key=lambda x: os.path.getctime(os.path.join(backup_dir, x)))
        
        return jsonify({
            'security_enabled': security_enabled,
            'integrity_check': integrity_ok,
            'database_stats': stats,
            'backup_info': {
                'count': backup_count,
                'latest': latest_backup,
                'directory': backup_dir
            },
            'environment': {
                'sqlite_secure': os.getenv('SQLITE_SECURE', 'false'),
                'backup_enabled': os.getenv('DB_BACKUP_ENABLED', 'false'),
                'encryption_key_set': bool(os.getenv('SQLITE_ENCRYPTION_KEY'))
            }
        })
        
    except Exception as e:
        current_app.logger.error(f"Erro ao obter status de segurança: {str(e)}")
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/security/backup', methods=['POST'])
@admin_required
def api_create_backup():
    """Criar backup do banco de dados"""
    try:
        from app import database as db
        
        backup_path = db.create_backup()
        
        if backup_path:
            return jsonify({
                'message': 'Backup criado com sucesso',
                'backup_path': backup_path,
                'timestamp': datetime.now().isoformat()
            })
        else:
            return jsonify({'error': 'Falha ao criar backup'}), 500
            
    except Exception as e:
        current_app.logger.error(f"Erro ao criar backup: {str(e)}")
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/security/integrity', methods=['POST'])
@admin_required
def api_check_integrity():
    """Verificar integridade do banco de dados"""
    try:
        from app import database as db
        
        integrity_ok = db.check_database_integrity()
        
        return jsonify({
            'integrity_check': integrity_ok,
            'message': 'Integridade OK' if integrity_ok else 'Problemas detectados na integridade',
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        current_app.logger.error(f"Erro na verificação de integridade: {str(e)}")
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/security/optimize', methods=['POST'])
@admin_required
def api_optimize_database():
    """Otimizar banco de dados"""
    try:
        from app import database as db
        
        success = db.optimize_database()
        
        return jsonify({
            'success': success,
            'message': 'Otimização concluída' if success else 'Falha na otimização',
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        current_app.logger.error(f"Erro na otimização: {str(e)}")
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/security/backups', methods=['GET'])
@admin_required
def api_list_backups():
    """Listar backups disponíveis"""
    try:
        from app import database as db
        
        backup_dir = os.path.join(os.path.dirname(db.DB_PATH), 'backups')
        backups = []
        
        if os.path.exists(backup_dir):
            for filename in os.listdir(backup_dir):
                if filename.endswith('.db'):
                    file_path = os.path.join(backup_dir, filename)
                    stat = os.stat(file_path)
                    
                    # Verificar se existe arquivo de hash
                    hash_file = f"{file_path}.sha256"
                    has_hash = os.path.exists(hash_file)
                    
                    backups.append({
                        'filename': filename,
                        'size': stat.st_size,
                        'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                        'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        'has_integrity_check': has_hash
                    })
        
        # Ordenar por data de criação (mais recente primeiro)
        backups.sort(key=lambda x: x['created_at'], reverse=True)
        
        return jsonify({
            'backups': backups,
            'backup_directory': backup_dir,
            'total_count': len(backups)
        })
        
    except Exception as e:
        current_app.logger.error(f"Erro ao listar backups: {str(e)}")
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/security/config', methods=['GET'])
@admin_required
def api_get_security_config():
    """Obter configuração de segurança atual"""
    try:
        config = {
            'environment_variables': {
                'SQLITE_SECURE': os.getenv('SQLITE_SECURE', 'false'),
                'DB_BACKUP_ENABLED': os.getenv('DB_BACKUP_ENABLED', 'false'),
                'SQLITE_ENCRYPTION_KEY': 'SET' if os.getenv('SQLITE_ENCRYPTION_KEY') else 'NOT_SET'
            },
            'recommendations': []
        }
        
        # Adicionar recomendações de segurança
        if os.getenv('SQLITE_SECURE', 'false').lower() != 'true':
            config['recommendations'].append({
                'type': 'warning',
                'message': 'SQLITE_SECURE não está habilitado. Configure para "true" para ativar configurações avançadas de segurança.'
            })
        
        if os.getenv('DB_BACKUP_ENABLED', 'false').lower() != 'true':
            config['recommendations'].append({
                'type': 'warning',
                'message': 'DB_BACKUP_ENABLED não está habilitado. Configure para "true" para ativar backups automáticos.'
            })
        
        if not os.getenv('SQLITE_ENCRYPTION_KEY'):
            config['recommendations'].append({
                'type': 'info',
                'message': 'SQLITE_ENCRYPTION_KEY não está configurada. Uma chave será gerada automaticamente, mas não será persistente entre reinicializações.'
            })
        
        return jsonify(config)
        
    except Exception as e:
        current_app.logger.error(f"Erro ao obter configuração de segurança: {str(e)}")
        return jsonify({'error': str(e)}), 500 

# ============ ROTAS DE GERENCIAMENTO DE API KEYS ============

@bp.route('/admin/api-keys')
@admin_required
def admin_api_keys():
    """Página de administração de API keys"""
    return render_template('admin/api_keys.html')

@bp.route('/admin/api/api-keys', methods=['GET'])
@admin_required
def api_list_api_keys():
    """Lista todas as API keys"""
    try:
        keys = db.get_all_api_keys()
        return jsonify({'api_keys': keys})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/api-keys', methods=['POST'])
@admin_required
def api_create_api_key():
    """Cria uma nova API key"""
    try:
        data = request.json
        name = data.get('name')
        description = data.get('description', '')
        permissions = data.get('permissions', 'read')
        expires_days = data.get('expires_days')
        
        if not name:
            return jsonify({'error': 'Nome da API key é obrigatório'}), 400
        
        # Validar permissões
        valid_permissions = ['read', 'write', 'admin']
        if permissions not in valid_permissions:
            return jsonify({'error': f'Permissões inválidas. Use: {", ".join(valid_permissions)}'}), 400
        
        # Obter user_id do usuário logado
        created_by = g.user.get('user_id') if hasattr(g, 'user') else None
        
        success, result = db.create_api_key(
            name=name,
            description=description,
            permissions=permissions,
            created_by=created_by,
            expires_days=expires_days
        )
        
        if success:
            return jsonify({
                'message': 'API key criada com sucesso',
                'api_key': result['api_key'],
                'key_id': result['id']
            }), 201
        else:
            return jsonify({'error': result}), 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/api-keys/<int:key_id>', methods=['DELETE'])
@admin_required
def api_delete_api_key(key_id):
    """Remove uma API key"""
    try:
        success, message = db.delete_api_key(key_id)
        
        if success:
            return jsonify({'message': message})
        else:
            return jsonify({'error': message}), 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/api-keys/<int:key_id>/toggle', methods=['PUT'])
@admin_required
def api_toggle_api_key(key_id):
    """Alternar ativação de uma chave de API"""
    try:
        # Obter o estado atual
        is_active = request.json.get('is_active', True)
        
        if db.toggle_api_key(key_id, is_active):
            return jsonify({
                'success': True,
                'message': f'Chave de API {"ativada" if is_active else "desativada"} com sucesso'
            })
        else:
            return jsonify({'error': 'Chave de API não encontrada'}), 404
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ==================== ROTAS PARA MULTAS E ADVERTÊNCIAS ====================



@bp.route('/admin/api/tipos-multa', methods=['GET'])
@admin_required
def api_get_tipos_multa():
    """Lista tipos de multa disponíveis"""
    try:
        tipos = db.get_tipos_multa()
        return jsonify(tipos)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/admin/api/tipos-advertencia', methods=['GET'])
@admin_required
def api_get_tipos_advertencia():
    """Lista tipos de advertência disponíveis"""
    try:
        tipos = db.get_tipos_advertencia()
        return jsonify(tipos)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/api/multas', methods=['GET'])
@login_required
def api_listar_multas():
    """Lista multas com filtros opcionais"""
    try:
        filtros = {}
        
        # Obter filtros da query string
        if request.args.get('unidade'):
            filtros['unidade'] = request.args.get('unidade')
        if request.args.get('status'):
            filtros['status'] = request.args.get('status')
        if request.args.get('data_inicio'):
            filtros['data_inicio'] = int(request.args.get('data_inicio'))
        if request.args.get('data_fim'):
            filtros['data_fim'] = int(request.args.get('data_fim'))
        
        multas = db.listar_multas(filtros if filtros else None)
        return jsonify(multas)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/api/multas', methods=['POST'])
@login_required
def api_criar_multa():
    """Cria uma nova multa"""
    try:
        data = request.json
        
        # Validar campos obrigatórios
        campos_obrigatorios = ['unidade', 'valor', 'data_infracao', 'descricao', 'assistant_id']
        for campo in campos_obrigatorios:
            if not data.get(campo):
                return jsonify({'error': f'Campo {campo} é obrigatório'}), 400
        
        # Adicionar usuário criador
        data['criado_por'] = g.user['user_id']
        
        # Converter data para timestamp
        if isinstance(data['data_infracao'], str):
            try:
                # Converter data string para objeto datetime
                data_obj = datetime.fromisoformat(data['data_infracao'].replace('Z', '+00:00'))
                
                # Validar se a data não é muito antiga (antes de 1900)
                if data_obj.year < 1900:
                    return jsonify({'error': 'Data da infração deve ser posterior a 1900'}), 400
                
                # Validar se a data não é muito no futuro (mais de 10 anos)
                data_limite = datetime.now().replace(year=datetime.now().year + 10)
                if data_obj > data_limite:
                    return jsonify({'error': 'Data da infração não pode ser mais de 10 anos no futuro'}), 400
                
                # Converter para timestamp
                timestamp = int(data_obj.timestamp())
                
                # Validar se o timestamp é positivo
                if timestamp < 0:
                    return jsonify({'error': 'Data da infração inválida - deve ser posterior a 1970'}), 400
                
                data['data_infracao'] = timestamp
                
            except ValueError as e:
                return jsonify({'error': f'Formato de data da infração inválido: {str(e)}'}), 400
        elif isinstance(data['data_infracao'], (int, float)):
            # Se já é um timestamp, validar se é válido
            timestamp = int(data['data_infracao'])
            
            # Validar se o timestamp é positivo
            if timestamp < 0:
                return jsonify({'error': 'Data da infração inválida - deve ser posterior a 1970'}), 400
            
            # Validar se não é muito no futuro (timestamp muito grande)
            limite_futuro = int((datetime.now().replace(year=datetime.now().year + 10)).timestamp())
            if timestamp > limite_futuro:
                return jsonify({'error': 'Data da infração não pode ser mais de 10 anos no futuro'}), 400
            
            # Validar se não é muito antiga (antes de 1900)
            limite_passado = int(datetime(1900, 1, 1).timestamp())
            if timestamp < limite_passado:
                return jsonify({'error': 'Data da infração deve ser posterior a 1900'}), 400
            
            data['data_infracao'] = timestamp
        
        # Se tiver data de vencimento, converter também
        if data.get('data_vencimento'):
            if isinstance(data['data_vencimento'], str):
                try:
                    # Converter data string para objeto datetime
                    data_obj = datetime.fromisoformat(data['data_vencimento'].replace('Z', '+00:00'))
                    
                    # Validar se a data não é muito antiga (antes de 1900)
                    if data_obj.year < 1900:
                        return jsonify({'error': 'Data de vencimento deve ser posterior a 1900'}), 400
                    
                    # Validar se a data não é muito no futuro (mais de 20 anos)
                    data_limite = datetime.now().replace(year=datetime.now().year + 20)
                    if data_obj > data_limite:
                        return jsonify({'error': 'Data de vencimento não pode ser mais de 20 anos no futuro'}), 400
                    
                    # Converter para timestamp
                    timestamp = int(data_obj.timestamp())
                    
                    # Validar se o timestamp é positivo
                    if timestamp < 0:
                        return jsonify({'error': 'Data de vencimento inválida - deve ser posterior a 1970'}), 400
                    
                    data['data_vencimento'] = timestamp
                    
                except ValueError as e:
                    return jsonify({'error': f'Formato de data de vencimento inválido: {str(e)}'}), 400
        
        # Criar multa
        multa = db.criar_multa(data)
        
        # Se tiver que gerar documento
        if data.get('gerar_documento'):
            # Preparar imagens se houver
            imagens_dados = []
            if data.get('imagens_info') and data['imagens_info'].get('imagens'):
                imagens_dados = data['imagens_info']['imagens']
            
            # Chamar função de geração de documento
            documento_data = {
                'tipo': 'multa',
                'dados': {
                    'numero': multa['numero_multa'],
                    'unidade': multa['unidade'],
                    'bloco': multa.get('bloco'),
                    'assistant_id': multa.get('assistant_id'),
                    'valor': multa['valor'],
                    'data': datetime.fromtimestamp(multa['data_infracao']).strftime('%d/%m/%Y'),
                    'descricao': multa['descricao'],
                    'observacoes': multa.get('observacoes', ''),
                    'documento_id': data.get('documento_id'),  # Para limpeza de imagens
                    'imagens': imagens_dados  # Imagens para inserir no documento
                }
            }
            
            current_app.logger.info(f"Gerando documento de multa com IA para assistant_id: {multa.get('assistant_id')}")
            
            response = gerar_documento_interno(documento_data)
            if response.get('documento_url'):
                # Atualizar multa com URL do documento
                conn = db.get_db_connection()
                cursor = conn.cursor()
                cursor.execute('UPDATE multas SET arquivo_documento = ? WHERE id = ?', 
                             (response['documento_url'], multa['id']))
                conn.commit()
                conn.close()
                multa['arquivo_documento'] = response['documento_url']
        
        return jsonify({
            'success': True,
            'multa': multa
        })
        
    except Exception as e:
        current_app.logger.error(f"Erro ao criar multa: {str(e)}")
        return jsonify({'error': str(e)}), 500

@bp.route('/api/advertencias', methods=['GET'])
@login_required
def api_listar_advertencias():
    """Lista advertências com filtros opcionais"""
    try:
        filtros = {}
        
        # Obter filtros da query string
        if request.args.get('unidade'):
            filtros['unidade'] = request.args.get('unidade')
        if request.args.get('status'):
            filtros['status'] = request.args.get('status')
        if request.args.get('data_inicio'):
            filtros['data_inicio'] = int(request.args.get('data_inicio'))
        if request.args.get('data_fim'):
            filtros['data_fim'] = int(request.args.get('data_fim'))
        
        advertencias = db.listar_advertencias(filtros if filtros else None)
        return jsonify(advertencias)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/api/advertencias', methods=['POST'])
@login_required
def api_criar_advertencia():
    """Cria uma nova advertência"""
    try:
        data = request.json
        
        # Validar campos obrigatórios
        campos_obrigatorios = ['unidade', 'data_ocorrencia', 'descricao', 'assistant_id']
        for campo in campos_obrigatorios:
            if not data.get(campo):
                return jsonify({'error': f'Campo {campo} é obrigatório'}), 400
        
        # Adicionar usuário criador
        data['criado_por'] = g.user['user_id']
        
        # Converter data para timestamp
        if isinstance(data['data_ocorrencia'], str):
            try:
                # Converter data string para objeto datetime
                data_obj = datetime.fromisoformat(data['data_ocorrencia'].replace('Z', '+00:00'))
                
                # Validar se a data não é muito antiga (antes de 1900)
                if data_obj.year < 1900:
                    return jsonify({'error': 'Data deve ser posterior a 1900'}), 400
                
                # Validar se a data não é muito no futuro (mais de 10 anos)
                data_limite = datetime.now().replace(year=datetime.now().year + 10)
                if data_obj > data_limite:
                    return jsonify({'error': 'Data não pode ser mais de 10 anos no futuro'}), 400
                
                # Converter para timestamp
                timestamp = int(data_obj.timestamp())
                
                # Validar se o timestamp é positivo
                if timestamp < 0:
                    return jsonify({'error': 'Data inválida - deve ser posterior a 1970'}), 400
                
                data['data_ocorrencia'] = timestamp
                
            except ValueError as e:
                return jsonify({'error': f'Formato de data inválido: {str(e)}'}), 400
        elif isinstance(data['data_ocorrencia'], (int, float)):
            # Se já é um timestamp, validar se é válido
            timestamp = int(data['data_ocorrencia'])
            
            # Validar se o timestamp é positivo
            if timestamp < 0:
                return jsonify({'error': 'Data inválida - deve ser posterior a 1970'}), 400
            
            # Validar se não é muito no futuro (timestamp muito grande)
            limite_futuro = int((datetime.now().replace(year=datetime.now().year + 10)).timestamp())
            if timestamp > limite_futuro:
                return jsonify({'error': 'Data não pode ser mais de 10 anos no futuro'}), 400
            
            # Validar se não é muito antiga (antes de 1900)
            limite_passado = int(datetime(1900, 1, 1).timestamp())
            if timestamp < limite_passado:
                return jsonify({'error': 'Data deve ser posterior a 1900'}), 400
            
            data['data_ocorrencia'] = timestamp
        
        # Criar advertência
        current_app.logger.info(f"Tentando criar advertência com dados: {data}")
        advertencia = db.criar_advertencia(data)
        current_app.logger.info(f"Advertência criada com sucesso: {advertencia}")
        
        # Se tiver que gerar documento
        if data.get('gerar_documento'):
            current_app.logger.info("Iniciando geração de documento...")
            # Preparar imagens se houver
            imagens_dados = []
            if data.get('imagens_info') and data['imagens_info'].get('imagens'):
                imagens_dados = data['imagens_info']['imagens']
            
            # Chamar função de geração de documento
            documento_data = {
                'tipo': 'advertencia',
                'dados': {
                    'numero': advertencia['numero_advertencia'],
                    'unidade': advertencia['unidade'],
                    'bloco': advertencia.get('bloco'),
                    'assistant_id': advertencia.get('assistant_id'),
                    'data': datetime.fromtimestamp(advertencia['data_ocorrencia']).strftime('%d/%m/%Y'),
                    'descricao': advertencia['descricao'],
                    'reincidente': advertencia.get('reincidente', 0),
                    'documento_id': data.get('documento_id'),  # Para limpeza de imagens
                    'imagens': imagens_dados  # Imagens para inserir no documento
                }
            }
            
            current_app.logger.info(f"Gerando documento de advertência com IA para assistant_id: {advertencia.get('assistant_id')}")
            
            response = gerar_documento_interno(documento_data)
            if response.get('documento_url'):
                # Atualizar advertência com URL do documento
                conn = db.get_db_connection()
                cursor = conn.cursor()
                cursor.execute('UPDATE advertencias SET arquivo_documento = ? WHERE id = ?', 
                             (response['documento_url'], advertencia['id']))
                conn.commit()
                conn.close()
                advertencia['arquivo_documento'] = response['documento_url']
        
        return jsonify({
            'success': True,
            'advertencia': advertencia
        }), 201
        
    except Exception as e:
        current_app.logger.error(f"Erro ao criar advertência: {str(e)}")
        return jsonify({'error': str(e)}), 500

@bp.route('/api/multas/<int:multa_id>', methods=['GET'])
@login_required
def api_get_multa(multa_id):
    """Obtém detalhes de uma multa específica"""
    try:
        multa = db.get_multa(multa_id)
        if multa:
            return jsonify(multa)
        return jsonify({'error': 'Multa não encontrada'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/api/advertencias/<int:advertencia_id>', methods=['GET'])
@login_required
def api_get_advertencia(advertencia_id):
    """Obtém detalhes de uma advertência específica"""
    try:
        advertencia = db.get_advertencia(advertencia_id)
        if advertencia:
            return jsonify(advertencia)
        return jsonify({'error': 'Advertência não encontrada'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/api/multas/<int:multa_id>/status', methods=['PUT'])
@login_required
def api_atualizar_status_multa(multa_id):
    """Atualiza o status de uma multa"""
    try:
        data = request.json
        status = data.get('status')
        
        if not status:
            return jsonify({'error': 'Status é obrigatório'}), 400
        
        if status not in ['pendente', 'paga', 'cancelada']:
            return jsonify({'error': 'Status inválido'}), 400
        
        data_pagamento = None
        if status == 'paga' and data.get('data_pagamento'):
            data_pagamento = int(datetime.fromisoformat(data['data_pagamento'].replace('Z', '+00:00')).timestamp())
        
        if db.atualizar_status_multa(multa_id, status, data_pagamento):
            return jsonify({
                'success': True,
                'message': f'Status da multa atualizado para {status}'
            })
        
        return jsonify({'error': 'Erro ao atualizar status'}), 500
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Função auxiliar para gerar documento internamente
def gerar_documento_interno(data):
    """Função interna para gerar documentos .docx com imagens"""
    tipo_documento = data.get('tipo')
    dados = data.get('dados')
    
    try:
        current_app.logger.info(f"gerar_documento_interno iniciado - tipo: {tipo_documento}")
        current_app.logger.info(f"Dados recebidos: {dados}")
        
        # Criar diretório para documentos se não existir
        docs_dir = os.path.join(current_app.root_path, 'static', 'documentos')
        current_app.logger.info(f"Diretório de documentos: {docs_dir}")
        if not os.path.exists(docs_dir):
            os.makedirs(docs_dir)
            current_app.logger.info("Diretório criado")
        
        # Gerar nome de arquivo único
        documento_id = str(uuid.uuid4())
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        current_app.logger.info(f"documento_id: {documento_id}, timestamp: {timestamp}")
        
        # Tentar gerar texto com IA se tiver assistant_id
        texto_gerado = None
        nome_assistant = None
        
        if dados.get('assistant_id'):
            current_app.logger.info(f"Tentando gerar texto com IA para assistant_id: {dados.get('assistant_id')}")
            try:
                # Buscar nome do assistant
                assistant = db.get_assistant_by_id(dados['assistant_id'])
                if assistant:
                    nome_assistant = assistant.get('name', 'Condomínio')
                current_app.logger.info(f"Nome do assistant: {nome_assistant}")
                
                # Preparar dados para IA
                ocorrencia = {
                    'external_assistant_id': dados['assistant_id'],
                    'morador': {
                        'nome': f"Morador(a) da unidade {dados.get('unidade', '')}",
                        'apartamento': dados.get('unidade', ''),
                        'bloco': dados.get('bloco', '')
                    },
                    'data': dados.get('data'),
                    'descricao': dados.get('descricao', ''),
                    'valor': float(dados.get('valor', 0)) if tipo_documento == 'multa' else None
                }
                current_app.logger.info(f"Dados para IA preparados: {ocorrencia}")
                
                # Importar e usar o serviço
                from app.services.openai_service import OpenAIService
                import asyncio
                
                service = OpenAIService()
                current_app.logger.info("OpenAIService criado")
                
                # Executar a função assíncrona
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                current_app.logger.info("Loop asyncio criado")
                
                texto_gerado = loop.run_until_complete(
                    service.gerar_documento_com_assistant(ocorrencia, tipo_documento)
                )
                current_app.logger.info(f"Texto gerado com IA: {texto_gerado[:100]}...")
                loop.close()
                
            except Exception as e:
                current_app.logger.error(f"Erro ao gerar com IA: {str(e)}")
                # Continuar com geração padrão
        
        # Gerar documento .docx
        current_app.logger.info(f"Gerando documento .docx do tipo: {tipo_documento}")
        if tipo_documento == 'advertencia':
            documento_url = gerar_advertencia_docx(
                documento_id, timestamp, dados, docs_dir, 
                texto_gerado, nome_assistant
            )
        elif tipo_documento == 'multa':
            documento_url = gerar_multa_docx(
                documento_id, timestamp, dados, docs_dir,
                texto_gerado, nome_assistant
            )
        else:
            raise ValueError('Tipo de documento não suportado')
        
        current_app.logger.info(f"Documento gerado: {documento_url}")
        
        # Limpar imagens temporárias após gerar o documento
        if dados.get('documento_id'):
            current_app.logger.info(f"Limpando imagens temporárias: {dados.get('documento_id')}")
            limpar_imagens_temporarias(dados['documento_id'])
        
        return {
            'success': True,
            'documento_url': documento_url
        }
        
    except Exception as e:
        current_app.logger.error(f"Erro ao gerar documento: {str(e)}")
        import traceback
        current_app.logger.error(f"Traceback: {traceback.format_exc()}")
        raise

# Função para gerar advertência em formato .docx
def gerar_advertencia_docx(documento_id, timestamp, dados, docs_dir, texto_gerado=None, nome_assistant=None):
    """Gera documento de advertência em formato .docx com imagens"""
    try:
        current_app.logger.info(f"gerar_advertencia_docx iniciado")
        current_app.logger.info(f"documento_id: {documento_id}")
        current_app.logger.info(f"timestamp: {timestamp}")
        current_app.logger.info(f"dados: {dados}")
        current_app.logger.info(f"docs_dir: {docs_dir}")
        
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_SECTION
        from docx.oxml.shared import OxmlElement, qn
        
        current_app.logger.info("Imports realizados com sucesso")
        
        # Criar documento
        doc = Document()
        current_app.logger.info("Document() criado")
        
        # Configurar margens
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Cabeçalho
        nome_condominio = nome_assistant or 'CONDOMÍNIO'
        header = doc.add_heading(nome_condominio, level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Título
        title = doc.add_heading('ADVERTÊNCIA DISCIPLINAR', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Número do documento
        numero_doc = doc.add_paragraph(f"Nº {dados.get('numero', 'S/N')}")
        numero_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Linha separadora
        doc.add_paragraph('_' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informações da unidade
        info_para = doc.add_paragraph()
        info_para.add_run('Unidade: ').bold = True
        info_para.add_run(f"{dados.get('unidade', 'Não informada')}")
        if dados.get('bloco'):
            info_para.add_run(f" - Bloco {dados.get('bloco')}")
        
        data_para = doc.add_paragraph()
        data_para.add_run('Data da Ocorrência: ').bold = True
        data_para.add_run(dados.get('data', 'Não informada'))
        
        # Espaço
        doc.add_paragraph()
        
        # PARTE 1: DESCRIÇÃO DO OCORRIDO
        if texto_gerado:
            # Dividir texto em descrição e fundamentação
            descricao, fundamentacao = dividir_texto_ia(texto_gerado)
            
            if descricao:
                # Adicionar descrição
                paragrafos_descricao = descricao.split('\n')
                for para_texto in paragrafos_descricao:
                    if para_texto.strip():
                        doc.add_paragraph(para_texto.strip())
        else:
            # Texto padrão - apenas a descrição
            doc.add_paragraph(f"Prezado Sr./Sra. Morador(a) da unidade {dados.get('unidade', 'Não informada')},")
            doc.add_paragraph(f"Informamos que foi constatada a seguinte ocorrência em {dados.get('data', 'Não informada')}:")
            doc.add_paragraph(dados.get('descricao', 'Não informada'))
            doc.add_paragraph("Esta atividade está sujeita a regulamentação específica.")
        
        # Espaço após descrição
        doc.add_paragraph()
        
        # INSERIR IMAGENS (entre descrição e fundamentação)
        if dados.get('imagens'):
            evidencias_para = doc.add_paragraph()
            evidencias_para.add_run('Evidências Fotográficas:').bold = True
            
            for imagem in dados['imagens']:
                try:
                    # Verificar se o arquivo existe
                    if os.path.exists(imagem['caminho']):
                        # Adicionar parágrafo para a imagem
                        img_para = doc.add_paragraph()
                        run = img_para.runs[0] if img_para.runs else img_para.add_run()
                        
                        # Inserir imagem com tamanho adequado
                        run.add_picture(imagem['caminho'], width=Inches(4))
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Adicionar legenda
                        caption = doc.add_paragraph(f"Imagem: {imagem['nome_original']}")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                except Exception as e:
                    current_app.logger.error(f"Erro ao inserir imagem {imagem['nome_original']}: {str(e)}")
                    # Continuar sem a imagem
            
            # Espaço após as imagens
            doc.add_paragraph()
        
        # PARTE 2: FUNDAMENTAÇÃO LEGAL
        if texto_gerado:
            _, fundamentacao = dividir_texto_ia(texto_gerado)
            
            if fundamentacao:
                # Adicionar fundamentação
                paragrafos_fundamentacao = fundamentacao.split('\n')
                for para_texto in paragrafos_fundamentacao:
                    if para_texto.strip():
                        doc.add_paragraph(para_texto.strip())
            else:
                # Se não tiver fundamentação específica, adicionar padrão
                doc.add_paragraph("Alertamos que, em caso de reincidência, serão aplicadas as sanções previstas no regulamento interno.")
        else:
            # Fundamentação padrão
            doc.add_paragraph("Alertamos que, em caso de reincidência, serão aplicadas as sanções previstas no regulamento interno.")
        
        # Finalização padrão - verificar se já tem "Atenciosamente"
        if texto_gerado:
            # Verificar se o texto já termina com atenciosamente
            texto_limpo = texto_gerado.strip().lower()
            if not texto_limpo.endswith('atenciosamente') and not texto_limpo.endswith('atenciosamente.'):
                doc.add_paragraph("Atenciosamente.")
        else:
            doc.add_paragraph("Atenciosamente.")
        
        # Espaço para assinaturas
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Tabela de assinaturas
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        # Linha 1 - espaço para assinatura
        table.cell(0, 0).text = ""
        table.cell(0, 1).text = ""
        
        # Linha 2 - labels
        table.cell(1, 0).text = "Administração do Condomínio"
        table.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 1).text = "Recebido por"
        table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Rodapé
        doc.add_paragraph()
        footer_para = doc.add_paragraph(f"Documento gerado eletronicamente em {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(10)
        
        # Salvar documento
        filename = f"advertencia_{timestamp}_{documento_id[:8]}.docx"
        filepath = os.path.join(docs_dir, filename)
        current_app.logger.info(f"Tentando salvar documento em: {filepath}")
        doc.save(filepath)
        current_app.logger.info(f"Documento salvo com sucesso: {filepath}")
        
        return f"/static/documentos/{filename}"
        
    except Exception as e:
        current_app.logger.error(f"Erro ao gerar advertência .docx: {str(e)}")
        raise

# Função para gerar multa em formato .docx
def gerar_multa_docx(documento_id, timestamp, dados, docs_dir, texto_gerado=None, nome_assistant=None):
    """Gera documento de multa em formato .docx com imagens"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor
        
        # Criar documento
        doc = Document()
        
        # Configurar margens
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Cabeçalho
        nome_condominio = nome_assistant or 'CONDOMÍNIO'
        header = doc.add_heading(nome_condominio, level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Título
        title = doc.add_heading('NOTIFICAÇÃO DE MULTA', level=2)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Número do documento
        numero_doc = doc.add_paragraph(f"Nº {dados.get('numero', 'S/N')}")
        numero_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Linha separadora
        doc.add_paragraph('_' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informações da unidade e valor em destaque
        info_para = doc.add_paragraph()
        info_para.add_run('Unidade: ').bold = True
        info_para.add_run(f"{dados.get('unidade', 'Não informada')}")
        if dados.get('bloco'):
            info_para.add_run(f" - Bloco {dados.get('bloco')}")
        
        data_para = doc.add_paragraph()
        data_para.add_run('Data da Infração: ').bold = True
        data_para.add_run(dados.get('data', 'Não informada'))
        
        # Valor da multa em destaque
        doc.add_paragraph()
        valor_para = doc.add_paragraph()
        valor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        valor_run = valor_para.add_run(f"VALOR DA MULTA: R$ {dados.get('valor', '0,00')}")
        valor_run.bold = True
        valor_run.font.size = Pt(16)
        valor_run.font.color.rgb = RGBColor(220, 38, 38)  # Vermelho
        
        # Espaço
        doc.add_paragraph()
        
        # PARTE 1: DESCRIÇÃO DO OCORRIDO
        if texto_gerado:
            # Dividir texto em descrição e fundamentação
            descricao, fundamentacao = dividir_texto_ia(texto_gerado)
            
            if descricao:
                # Adicionar descrição
                paragrafos_descricao = descricao.split('\n')
                for para_texto in paragrafos_descricao:
                    if para_texto.strip():
                        doc.add_paragraph(para_texto.strip())
        else:
            # Texto padrão - apenas a descrição
            valor_formatado = dados.get('valor', '0,00')
            doc.add_paragraph(f"Prezado Sr./Sra. Morador(a) da unidade {dados.get('unidade', 'Não informada')},")
            doc.add_paragraph(f"Informamos que foi aplicada multa no valor de R$ {valor_formatado} referente à infração constatada em {dados.get('data', 'Não informada')}:")
            doc.add_paragraph(dados.get('descricao', 'Não informada'))
            doc.add_paragraph("Esta infração está sujeita à penalidade financeira.")
        
        # Espaço após descrição
        doc.add_paragraph()
        
        # INSERIR IMAGENS (entre descrição e fundamentação)
        if dados.get('imagens'):
            evidencias_para = doc.add_paragraph()
            evidencias_para.add_run('Evidências Fotográficas:').bold = True
            
            for imagem in dados['imagens']:
                try:
                    # Verificar se o arquivo existe
                    if os.path.exists(imagem['caminho']):
                        # Adicionar parágrafo para a imagem
                        img_para = doc.add_paragraph()
                        run = img_para.runs[0] if img_para.runs else img_para.add_run()
                        
                        # Inserir imagem com tamanho adequado
                        run.add_picture(imagem['caminho'], width=Inches(4))
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Adicionar legenda
                        caption = doc.add_paragraph(f"Imagem: {imagem['nome_original']}")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                except Exception as e:
                    current_app.logger.error(f"Erro ao inserir imagem {imagem['nome_original']}: {str(e)}")
                    # Continuar sem a imagem
            
            # Espaço após as imagens
            doc.add_paragraph()
        
        # PARTE 2: FUNDAMENTAÇÃO LEGAL
        if texto_gerado:
            _, fundamentacao = dividir_texto_ia(texto_gerado)
            
            if fundamentacao:
                # Adicionar fundamentação
                paragrafos_fundamentacao = fundamentacao.split('\n')
                for para_texto in paragrafos_fundamentacao:
                    if para_texto.strip():
                        doc.add_paragraph(para_texto.strip())
            else:
                # Se não tiver fundamentação específica, adicionar padrão
                doc.add_paragraph("Conforme estabelecido no regimento interno do condomínio.")
                doc.add_paragraph("O pagamento deverá ser realizado junto à administração do condomínio no prazo de 30 dias a contar desta notificação.")
        else:
            # Fundamentação padrão
            doc.add_paragraph("Conforme estabelecido no regimento interno do condomínio, esta infração está sujeita à penalidade financeira.")
            doc.add_paragraph("O pagamento deverá ser realizado junto à administração do condomínio no prazo de 30 dias a contar desta notificação.")
        
        # Finalização padrão - verificar se já tem "Atenciosamente"
        if texto_gerado:
            # Verificar se o texto já termina com atenciosamente
            texto_limpo = texto_gerado.strip().lower()
            if not texto_limpo.endswith('atenciosamente') and not texto_limpo.endswith('atenciosamente.'):
                doc.add_paragraph("Atenciosamente.")
        else:
            doc.add_paragraph("Atenciosamente.")
        
        # Espaço para assinaturas
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Tabela de assinaturas
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        # Linha 1 - espaço para assinatura
        table.cell(0, 0).text = ""
        table.cell(0, 1).text = ""
        
        # Linha 2 - labels
        table.cell(1, 0).text = "Administração do Condomínio"
        table.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, 1).text = "Recebido por"
        table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Rodapé
        doc.add_paragraph()
        footer_para = doc.add_paragraph(f"Documento gerado eletronicamente em {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(10)
        
        # Salvar documento
        filename = f"multa_{timestamp}_{documento_id[:8]}.docx"
        filepath = os.path.join(docs_dir, filename)
        doc.save(filepath)
        
        return f"/static/documentos/{filename}"
        
    except Exception as e:
        current_app.logger.error(f"Erro ao gerar multa .docx: {str(e)}")
        raise

# Função para limpar imagens temporárias
def limpar_imagens_temporarias(documento_id):
    """Remove imagens temporárias após gerar o documento"""
    try:
        temp_dir = os.path.join(current_app.root_path, 'static', 'temp_images', documento_id)
        if os.path.exists(temp_dir):
            import shutil
            shutil.rmtree(temp_dir)
            current_app.logger.info(f"Imagens temporárias removidas: {documento_id}")
    except Exception as e:
        current_app.logger.error(f"Erro ao limpar imagens temporárias: {str(e)}")

@bp.route('/upload-imagens-documento', methods=['POST'])
@login_required
def upload_imagens_documento():
    """Upload de imagens temporárias para documentos"""
    try:
        documento_id = request.form.get('documento_id')
        tipo_documento = request.form.get('tipo')
        
        if not documento_id or not tipo_documento:
            return jsonify({'error': 'ID do documento e tipo são obrigatórios'}), 400
        
        # Criar diretório temporário para imagens
        temp_dir = os.path.join(current_app.root_path, 'static', 'temp_images', documento_id)
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)
        
        imagens_salvas = []
        
        # Processar cada imagem enviada
        for key in request.files:
            if key.startswith('imagem_'):
                file = request.files[key]
                
                if file and file.filename:
                    # Validar formato
                    if not file.filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                        return jsonify({'error': f'Formato não suportado: {file.filename}'}), 400
                    
                    # Gerar nome único para a imagem
                    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
                    ext = os.path.splitext(file.filename)[1]
                    nome_arquivo = f"img_{len(imagens_salvas)}_{timestamp}{ext}"
                    
                    # Salvar arquivo
                    caminho_arquivo = os.path.join(temp_dir, nome_arquivo)
                    file.save(caminho_arquivo)
                    
                    # Adicionar à lista
                    imagens_salvas.append({
                        'nome_original': file.filename,
                        'nome_arquivo': nome_arquivo,
                        'caminho': caminho_arquivo,
                        'url_relativa': f"/static/temp_images/{documento_id}/{nome_arquivo}"
                    })
        
        if not imagens_salvas:
            return jsonify({'error': 'Nenhuma imagem válida foi enviada'}), 400
        
        return jsonify({
            'success': True,
            'imagens': imagens_salvas,
            'total': len(imagens_salvas)
        })
        
    except Exception as e:
        current_app.logger.error(f"Erro no upload de imagens: {str(e)}")
        return jsonify({'error': str(e)}), 500 

@bp.route('/admin/limpar-imagens-antigas', methods=['POST'])
@admin_required
def limpar_imagens_antigas():
    """Limpa imagens temporárias antigas (mais de 24 horas)"""
    try:
        import time
        import shutil
        
        temp_base_dir = os.path.join(current_app.root_path, 'static', 'temp_images')
        if not os.path.exists(temp_base_dir):
            return jsonify({'message': 'Diretório de imagens temporárias não encontrado'})
        
        agora = time.time()
        limite_tempo = 24 * 60 * 60  # 24 horas em segundos
        removidos = 0
        
        for pasta in os.listdir(temp_base_dir):
            pasta_path = os.path.join(temp_base_dir, pasta)
            if os.path.isdir(pasta_path):
                # Verificar idade da pasta
                idade = agora - os.path.getctime(pasta_path)
                if idade > limite_tempo:
                    try:
                        shutil.rmtree(pasta_path)
                        removidos += 1
                        current_app.logger.info(f"Pasta temporária antiga removida: {pasta}")
                    except Exception as e:
                        current_app.logger.error(f"Erro ao remover pasta {pasta}: {str(e)}")
        
        return jsonify({
            'success': True,
            'message': f'{removidos} pasta(s) de imagens antigas removida(s)',
            'removidos': removidos
        })
        
    except Exception as e:
        current_app.logger.error(f"Erro na limpeza de imagens antigas: {str(e)}")
        return jsonify({'error': str(e)}), 500 

# Função auxiliar para dividir texto da IA em descrição e fundamentação
def dividir_texto_ia(texto_gerado):
    """
    Divide o texto gerado pela IA em duas partes usando a API da OpenAI:
    1. Descrição do ocorrido (até antes da fundamentação legal)
    2. Fundamentação legal (artigos, regimentos, etc.)
    """
    if not texto_gerado:
        return None, None
    
    try:
        # Usar a API da OpenAI para dividir o texto de forma inteligente
        from app.services.openai_service import OpenAIService
        import asyncio
        
        service = OpenAIService()
        
        # Se não tiver cliente configurado, usar fallback regex
        if not service.client:
            current_app.logger.warning("Cliente OpenAI não disponível, usando fallback regex")
            return dividir_texto_fallback(texto_gerado)
        
        # Prompt para dividir o texto
        prompt = f"""Você recebeu o seguinte texto de um documento de advertência/multa:

{texto_gerado}

TAREFA: Divida este texto em EXATAMENTE duas partes para que possamos inserir imagens fotográficas entre elas:

1. **PARTE 1** - DESCRIÇÃO: Desde o início até o final da descrição da ocorrência/incidente
2. **PARTE 2** - FUNDAMENTAÇÃO: A partir da fundamentação legal, artigos, regimentos, consequências, etc.

As imagens devem ficar APÓS a descrição da ocorrência e ANTES da fundamentação legal.

IMPORTANTE: Retorne EXATAMENTE no formato solicitado, sem texto adicional:

===DESCRIÇÃO===
[todo o texto da descrição da ocorrência]

===FUNDAMENTAÇÃO===
[todo o texto de fundamentação legal e consequências]"""

        # Fazer chamada para a API
        response = service.client.chat.completions.create(
            model="gpt-4",
            messages=[
                {
                    "role": "system",
                    "content": "Você é um assistente especializado em dividir textos de documentos legais. Siga exatamente o formato solicitado."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.1,  # Baixa temperatura para consistência
            max_tokens=2000
        )
        
        resultado = response.choices[0].message.content
        
        # Fazer parsing da resposta
        if "===DESCRIÇÃO===" in resultado and "===FUNDAMENTAÇÃO===" in resultado:
            partes = resultado.split("===DESCRIÇÃO===")
            if len(partes) > 1:
                resto = partes[1]
                if "===FUNDAMENTAÇÃO===" in resto:
                    descricao_e_fundamentacao = resto.split("===FUNDAMENTAÇÃO===")
                    descricao = descricao_e_fundamentacao[0].strip()
                    fundamentacao = descricao_e_fundamentacao[1].strip() if len(descricao_e_fundamentacao) > 1 else ""
                    
                    # Log do resultado
                    current_app.logger.info("Texto dividido com sucesso pela API da OpenAI")
                    current_app.logger.info(f"Descrição: {descricao[:100]}...")
                    current_app.logger.info(f"Fundamentação: {fundamentacao[:100]}...")
                    
                    return descricao, fundamentacao
        
        # Se não conseguiu fazer parsing, usar fallback
        current_app.logger.warning("Falha no parsing da resposta da API, usando fallback")
        return dividir_texto_fallback(texto_gerado)
        
    except Exception as e:
        current_app.logger.error(f"Erro ao dividir texto com OpenAI: {str(e)}")
        # Fallback para método regex em caso de erro
        return dividir_texto_fallback(texto_gerado)

def dividir_texto_fallback(texto_gerado):
    """
    Função de fallback usando padrões regex (método anterior)
    """
    if not texto_gerado:
        return None, None
    
    # Padrões que indicam início da fundamentação legal
    padroes_fundamentacao = [
        r'De acordo com (?:o|a|os|as)',
        r'Conforme (?:o|a|os|as)',
        r'Segundo (?:o|a|os|as)',
        r'Baseado no?',
        r'Nos termos do?',
        r'Art\.?\s*\d',
        r'Artigo\s*\d',
        r'Lei n[º°]?\s*\d',
        r'Regimento Interno',
        r'Convenção do Condomínio',
        r'Lei de Condomínio',
        r'É defeso a qualquer condômino',
        r'É vedado',
        r'É proibido',
        r'alertamos que',
        r'em caso de reincidência',
        r'serão aplicadas sanções'
    ]
    
    # Procurar pela primeira ocorrência de fundamentação legal
    linhas = texto_gerado.split('\n')
    indice_fundamentacao = -1
    
    for i, linha in enumerate(linhas):
        linha_limpa = linha.strip()
        if linha_limpa:
            for padrao in padroes_fundamentacao:
                import re
                if re.search(padrao, linha_limpa, re.IGNORECASE):
                    indice_fundamentacao = i
                    break
            if indice_fundamentacao != -1:
                break
    
    if indice_fundamentacao != -1:
        # Dividir o texto
        descricao_linhas = linhas[:indice_fundamentacao]
        fundamentacao_linhas = linhas[indice_fundamentacao:]
        
        descricao = '\n'.join(descricao_linhas).strip()
        fundamentacao = '\n'.join(fundamentacao_linhas).strip()
        
        current_app.logger.info(f"Texto dividido com fallback regex - linha {indice_fundamentacao}")
        
        return descricao, fundamentacao
    else:
        # Se não encontrou fundamentação específica, dividir na metade
        meio = len(linhas) // 2
        descricao_linhas = linhas[:meio]
        fundamentacao_linhas = linhas[meio:]
        
        descricao = '\n'.join(descricao_linhas).strip()
        fundamentacao = '\n'.join(fundamentacao_linhas).strip()
        
        current_app.logger.info("Fundamentação não detectada automaticamente, dividindo na metade")
        
        return descricao, fundamentacao

# ==================== ROTAS PARA TRANSCRIÇÃO DE ÁUDIO ====================

@bp.route('/admin/transcricao')
@login_required  # Mudança: permitir todos os usuários logados, não apenas admins
def admin_transcricao():
    """Página principal de transcrição de áudio - Disponível para todos os usuários"""
    return render_template('admin/transcricao.html', user=g.user)

@bp.route('/api/transcricao/upload', methods=['POST'])
@login_required  # Mudança: permitir todos os usuários logados
def api_transcricao_upload():
    """Upload e processamento direto da transcrição de áudio usando microserviço transcrever"""
    try:
        # Verificar se o serviço de transcrição está disponível
        if not transcriber_client.health_check():
            return jsonify({
                'success': False, 
                'error': 'Serviço de transcrição indisponível. Verifique se o microserviço está rodando.'
            }), 503
        
        if 'audio' not in request.files:
            return jsonify({'success': False, 'error': 'Nenhum arquivo enviado'}), 400
        
        file = request.files['audio']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'Nenhum arquivo selecionado'}), 400
        
        # Parâmetros opcionais
        language = request.form.get('language', 'pt')
        speaker_labels = request.form.get('speaker_labels', 'false').lower() == 'true'
        
        # Gerar ID único para esta transcrição
        transcription_id = str(uuid.uuid4())
        filename = secure_filename(file.filename)
        
        # Salvar arquivo temporariamente
        temp_dir = os.path.join(current_app.root_path, 'static', 'temp_audio')
        os.makedirs(temp_dir, exist_ok=True)
        
        temp_filename = f"{transcription_id}_{filename}"
        file_path = os.path.join(temp_dir, temp_filename)
        file.save(file_path)
        
        current_app.logger.info(f"Arquivo salvo temporariamente: {file_path}")
        current_app.logger.info(f"Enviando arquivo para transcrição: {file_path}")
        
        # Fazer upload para o microserviço transcrever
        upload_result = transcriber_client.upload_audio_file(file_path, language, speaker_labels)
        if not upload_result or not upload_result.get('success'):
            # Limpar arquivo temporário
            try:
                os.remove(file_path)
            except:
                pass
            return jsonify({
                'success': False,
                'error': 'Erro ao enviar arquivo para transcrição'
            }), 500
        
        task_id = upload_result.get('task_id')
        current_app.logger.info(f"Upload realizado com sucesso. Task ID: {task_id}")
        current_app.logger.info(f"Task ID obtido: {task_id}")
        
        # Aguardar processamento com polling
        max_wait_time = 3600  # 60 minutos (1 hora) para arquivos grandes
        check_interval = 3   # 3 segundos
        start_time = time.time()
        
        while time.time() - start_time < max_wait_time:
            # Verificar status no transcrever
            status_response = transcriber_client.check_task_status(task_id)
            
            if not status_response:
                time.sleep(check_interval)
                continue
                
            state = status_response.get('state', 'UNKNOWN')
            current_app.logger.info(f"Task {task_id} status: {state}")
            
            if state == 'SUCCESS':
                # Sucesso! Baixar o resultado
                result = status_response.get('result', {})
                docx_filename = result.get('docx')
                
                if docx_filename:
                    # Baixar arquivo do transcrever
                    docx_content = transcriber_client.download_file(docx_filename)
                    
                    if docx_content:
                        current_app.logger.info(f"Transcrição {transcription_id} concluída com sucesso!")
                        
                        # Limpar arquivo temporário
                        try:
                            os.remove(file_path)
                        except:
                            pass
                        
                        # Retornar o arquivo diretamente para download
                        from flask import Response
                        
                        def generate():
                            yield docx_content
                        
                        return Response(
                            generate(),
                            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                            headers={
                                'Content-Disposition': f'attachment; filename="transcricao_{filename}.docx"',
                                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                            }
                        )
                    else:
                        return jsonify({
                            'success': False,
                            'error': 'Erro ao baixar resultado do serviço de transcrição'
                        }), 500
                else:
                    return jsonify({
                        'success': False,
                        'error': 'Resultado não disponível no serviço de transcrição'
                    }), 500
                
            elif state == 'FAILURE':
                error_msg = status_response.get('error', 'Erro desconhecido no processamento')
                current_app.logger.error(f"Transcrição {transcription_id} falhou: {error_msg}")
                
                # Limpar arquivo temporário
                try:
                    os.remove(file_path)
                except:
                    pass
                
                return jsonify({
                    'success': False,
                    'error': f'Erro no processamento: {error_msg}'
                }), 500
            
            # Aguardar antes da próxima verificação
            time.sleep(check_interval)
        
        # Timeout
        current_app.logger.warning(f"Timeout no processamento da transcrição {transcription_id}")
        
        # Limpar arquivo temporário
        try:
            os.remove(file_path)
        except:
            pass
        
        return jsonify({
            'success': False,
            'error': f'Timeout após {max_wait_time} segundos. Tente novamente com um arquivo menor.'
        }), 408
        
    except Exception as e:
        current_app.logger.error(f"Erro no upload de transcrição: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@bp.route('/api/transcricao/status/<transcription_id>')
@login_required  # Mudança: permitir todos os usuários logados
def api_transcricao_status(transcription_id):
    """Verifica o status da transcrição"""
    try:
        transcricao = db.get_transcricao_status(transcription_id)
        
        if not transcricao:
            return jsonify({'error': 'Transcrição não encontrada'}), 404
        
        response_data = {
            'status': transcricao['status'],
            'progress': transcricao.get('progress', 0),
            'message': transcricao.get('message', ''),
        }
        
        if transcricao['status'] == 'completed':
            response_data.update({
                'text': transcricao['texto'],
                'download_url': f"/api/transcricao/download/{transcription_id}"
            })
        elif transcricao['status'] == 'failed':
            response_data['error'] = transcricao.get('error_message', 'Erro desconhecido')
        
        return jsonify(response_data)
        
    except Exception as e:
        current_app.logger.error(f"Erro ao verificar status: {str(e)}")
        return jsonify({'error': str(e)}), 500

@bp.route('/api/transcricao/download/<transcription_id>')
@login_required  # Mudança: permitir todos os usuários logados
def api_transcricao_download(transcription_id):
    """Download do documento Word com a transcrição"""
    try:
        transcricao = db.get_transcricao_status(transcription_id)
        
        if not transcricao or transcricao['status'] != 'completed':
            return jsonify({'error': 'Transcrição não encontrada ou não concluída'}), 404
        
        word_path = transcricao.get('arquivo_word')
        if not word_path or not os.path.exists(word_path):
            return jsonify({'error': 'Arquivo Word não encontrado'}), 404
        
        return send_file(
            word_path,
            as_attachment=True,
            download_name=f"transcricao_{transcricao['nome_arquivo']}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        current_app.logger.error(f"Erro no download: {str(e)}")
        return jsonify({'error': str(e)}), 500

@bp.route('/api/transcricao/stats')
@login_required  # Mudança: permitir todos os usuários logados, mas mostrar apenas estatísticas próprias para usuários comuns
def api_transcricao_stats():
    """Retorna estatísticas de transcrições"""
    try:
        stats = db.get_transcricao_stats()
        return jsonify(stats)
    except Exception as e:
        current_app.logger.error(f"Erro ao carregar estatísticas: {str(e)}")
        return jsonify({'error': str(e)}), 500

@bp.route('/api/transcricao/recent')
@login_required  # Mudança: permitir todos os usuários logados, mas mostrar apenas transcrições próprias para usuários comuns
def api_transcricao_recent():
    """Retorna transcrições recentes"""
    try:
        recent = db.get_transcricoes_recentes(limit=5)
        return jsonify({'files': recent})
    except Exception as e:
        current_app.logger.error(f"Erro ao carregar arquivos recentes: {str(e)}")
        return jsonify({'error': str(e)}), 500

def process_transcription(transcription_id, file_path, language, speaker_labels):
    """Processa a transcrição de áudio usando AssemblyAI (simulado)"""
    try:
        # Atualizar status para processando
        db.update_transcricao_status(transcription_id, 'processing', 10, 'Iniciando processamento...')
        time.sleep(2)
        
        # Simular processamento (em produção, usar AssemblyAI API)
        db.update_transcricao_status(transcription_id, 'processing', 30, 'Enviando para AssemblyAI...')
        time.sleep(3)
        
        db.update_transcricao_status(transcription_id, 'processing', 60, 'Transcrevendo áudio...')
        time.sleep(5)
        
        # Texto de exemplo (substituir por chamada real da API)
        exemplo_texto = """
[SPEAKER 1]: Bom dia, como posso ajudá-lo hoje?

[SPEAKER 2]: Olá, eu gostaria de saber sobre os serviços disponíveis.

[SPEAKER 1]: Claro! Temos diversos serviços disponíveis. Você tem algum interesse específico?

[SPEAKER 2]: Sim, estou interessado em consultoria jurídica.

[SPEAKER 1]: Perfeito! Nossa equipe de advogados pode auxiliá-lo com diversos tipos de questões jurídicas.
        """.strip()
        
        db.update_transcricao_status(transcription_id, 'processing', 80, 'Gerando documento Word...')
        time.sleep(2)
        
        # Criar documento Word
        word_path = create_transcription_document(transcription_id, exemplo_texto, file_path)
        
        # Atualizar como concluído
        db.update_transcricao_status(
            transcription_id, 
            'completed', 
            100, 
            'Transcrição concluída!',
            texto=exemplo_texto,
            arquivo_word=word_path
        )
        
        # Limpar arquivo temporário
        try:
            os.remove(file_path)
        except:
            pass
            
    except Exception as e:
        current_app.logger.error(f"Erro no processamento da transcrição {transcription_id}: {str(e)}")
        db.update_transcricao_status(
            transcription_id, 
            'failed', 
            0, 
            f'Erro no processamento: {str(e)}'
        )

def create_transcription_document(transcription_id, text, original_file_path):
    """Cria documento Word com a transcrição"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        # Criar documento
        doc = Document()
        
        # Título
        title = doc.add_heading('Transcrição de Áudio', 0)
        title.alignment = 1  # Centralizado
        
        # Informações
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        
        info_table.cell(0, 0).text = 'Data da Transcrição:'
        info_table.cell(0, 1).text = datetime.now().strftime('%d/%m/%Y %H:%M')
        
        info_table.cell(1, 0).text = 'Arquivo Original:'
        info_table.cell(1, 1).text = os.path.basename(original_file_path)
        
        info_table.cell(2, 0).text = 'ID da Transcrição:'
        info_table.cell(2, 1).text = transcription_id
        
        # Adicionar quebra de página
        doc.add_page_break()
        
        # Título da transcrição
        doc.add_heading('Transcrição', level=1)
        
        # Adicionar texto da transcrição
        paragraphs = text.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Salvar documento
        output_dir = os.path.join(current_app.root_path, 'static', 'transcricoes')
        os.makedirs(output_dir, exist_ok=True)
        
        word_filename = f"transcricao_{transcription_id}.docx"
        word_path = os.path.join(output_dir, word_filename)
        
        doc.save(word_path)
        return word_path
        
    except Exception as e:
        current_app.logger.error(f"Erro ao criar documento Word: {str(e)}")
        raise e

# ... existing code ...
        current_app.logger.info("Fundamentação não detectada automaticamente, dividindo na metade")
        
        return descricao, fundamentacao

@bp.route('/documentos')
@login_required
def documentos():
    """Página principal de gestão de documentos (multas e advertências) - Interface única para todos"""
    return render_template('admin/documentos.html', user=g.user)

@bp.route('/api/documentos/estatisticas', methods=['GET'])
@login_required
def api_documentos_estatisticas():
    """Retorna estatísticas de multas e advertências"""
    try:
        stats = db.get_estatisticas_documentos()
        return jsonify(stats)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def monitor_transcription(transcription_id, transcrever_task_id, local_file_path):
    """Monitora o progresso da transcrição no microserviço transcrever"""
    try:
        current_app.logger.info(f"Iniciando monitoramento da transcrição {transcription_id} (task: {transcrever_task_id})")
        
        # Monitorar progresso
        max_wait_time = 7200  # 2 horas máximo
        start_time = time.time()
        check_interval = 10  # Verificar a cada 10 segundos
        
        while time.time() - start_time < max_wait_time:
            # Verificar status no transcrever
            status_response = transcriber_client.get_task_status(transcrever_task_id)
            
            if not status_response:
                current_app.logger.error(f"Falha ao obter status da task {transcrever_task_id}")
                db.update_transcricao_status(
                    transcription_id, 
                    'failed', 
                    0, 
                    'Falha na comunicação com o serviço de transcrição'
                )
                break
            
            state = status_response.get('state', '').upper()
            current_app.logger.debug(f"Task {transcrever_task_id} status: {state}")
            
            if state == 'PENDING':
                db.update_transcricao_status(transcription_id, 'processing', 10, 'Na fila de processamento...')
                
            elif state == 'PROGRESS':
                # Se tiver informação de progresso específica
                status_msg = status_response.get('status', 'Processando...')
                db.update_transcricao_status(transcription_id, 'processing', 50, status_msg)
                
            elif state == 'SUCCESS':
                current_app.logger.info(f"Transcrição {transcription_id} concluída com sucesso!")
                
                # Obter resultado
                result = status_response.get('result', {})
                docx_filename = result.get('docx')
                
                if docx_filename:
                    # Baixar arquivo do transcrever
                    local_transcricoes_dir = os.path.join(current_app.root_path, 'static', 'transcricoes')
                    os.makedirs(local_transcricoes_dir, exist_ok=True)
                    
                    local_docx_path = os.path.join(local_transcricoes_dir, f"transcricao_{transcription_id}.docx")
                    
                    download_success = transcriber_client.download_result(docx_filename, local_docx_path)
                    
                    if download_success:
                        # Extrair texto do documento para salvar no banco
                        texto_extraido = extract_text_from_docx(local_docx_path)
                        
                        # Atualizar como concluído
                        db.update_transcricao_status(
                            transcription_id,
                            'completed',
                            100,
                            'Transcrição concluída com sucesso!',
                            texto=texto_extraido,
                            arquivo_word=local_docx_path
                        )
                        
                        current_app.logger.info(f"Transcrição {transcription_id} salva localmente: {local_docx_path}")
                    else:
                        db.update_transcricao_status(
                            transcription_id,
                            'failed',
                            90,
                            'Erro ao baixar resultado do serviço de transcrição'
                        )
                else:
                    db.update_transcricao_status(
                        transcription_id,
                        'failed',
                        90,
                        'Resultado não disponível no serviço de transcrição'
                    )
                break
                
            elif state == 'FAILURE':
                error_msg = status_response.get('error', 'Erro desconhecido no processamento')
                current_app.logger.error(f"Transcrição {transcription_id} falhou: {error_msg}")
                
                db.update_transcricao_status(
                    transcription_id,
                    'failed',
                    0,
                    f'Erro no processamento: {error_msg}'
                )
                break
                
            else:
                # Estado desconhecido, continuar monitorando
                current_app.logger.warning(f"Estado desconhecido para task {transcrever_task_id}: {state}")
            
            time.sleep(check_interval)
        
        else:
            # Timeout
            current_app.logger.warning(f"Timeout no monitoramento da transcrição {transcription_id}")
            db.update_transcricao_status(
                transcription_id,
                'failed',
                0,
                f'Timeout após {max_wait_time} segundos'
            )
        
        # Limpar arquivo temporário local
        try:
            if os.path.exists(local_file_path):
                os.remove(local_file_path)
                current_app.logger.info(f"Arquivo temporário removido: {local_file_path}")
        except Exception as e:
            current_app.logger.warning(f"Erro ao remover arquivo temporário: {str(e)}")
            
    except Exception as e:
        current_app.logger.error(f"Erro no monitoramento da transcrição {transcription_id}: {str(e)}")
        db.update_transcricao_status(
            transcription_id,
            'failed',
            0,
            f'Erro no monitoramento: {str(e)}'
        )

def extract_text_from_docx(docx_path):
    """Extrai texto de um arquivo DOCX para salvar no banco de dados"""
    try:
        from docx import Document
        
        doc = Document(docx_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        text = '\n'.join(full_text)
        return text[:5000]  # Limitar a 5000 caracteres para o banco
        
    except Exception as e:
        current_app.logger.error(f"Erro ao extrair texto do DOCX: {str(e)}")
        return "Erro ao extrair texto do documento"

@bp.route('/api/transcricao/files', methods=['GET'])
@login_required
def api_transcricao_files():
    """Lista arquivos de transcrição disponíveis para download"""
    try:
        transcricoes_dir = os.path.join(current_app.root_path, 'static', 'transcricoes')
        
        if not os.path.exists(transcricoes_dir):
            return jsonify({'files': []})
        
        files = []
        for filename in os.listdir(transcricoes_dir):
            if filename.endswith('.docx'):
                file_path = os.path.join(transcricoes_dir, filename)
                if os.path.isfile(file_path):
                    # Obter informações do arquivo
                    stat = os.stat(file_path)
                    size = stat.st_size
                    modified = datetime.fromtimestamp(stat.st_mtime)
                    
                    # Extrair nome original do arquivo
                    original_name = filename
                    if filename.startswith('transcricao_') and '_' in filename:
                        # Formato: transcricao_UUID_original_name.docx
                        parts = filename.split('_', 2)
                        if len(parts) >= 3:
                            original_name = parts[2]
                    
                    files.append({
                        'filename': filename,
                        'original_name': original_name,
                        'size': size,
                        'size_formatted': format_file_size(size),
                        'modified': modified.strftime('%d/%m/%Y %H:%M'),
                        'download_url': f'/api/transcricao/download-file/{filename}'
                    })
        
        # Ordenar por data de modificação (mais recente primeiro)
        files.sort(key=lambda x: x['modified'], reverse=True)
        
        return jsonify({'files': files})
        
    except Exception as e:
        current_app.logger.error(f"Erro ao listar arquivos: {str(e)}")
        return jsonify({'error': str(e)}), 500

@bp.route('/api/transcricao/download-file/<filename>')
@login_required
def api_transcricao_download_file(filename):
    """Download de arquivo específico de transcrição"""
    try:
        # Validar nome do arquivo (segurança)
        if not filename.endswith('.docx') or '..' in filename or '/' in filename:
            return jsonify({'error': 'Nome de arquivo inválido'}), 400
        
        transcricoes_dir = os.path.join(current_app.root_path, 'static', 'transcricoes')
        file_path = os.path.join(transcricoes_dir, filename)
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'Arquivo não encontrado'}), 404
        
        return send_file(
            file_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        current_app.logger.error(f"Erro no download do arquivo: {str(e)}")
        return jsonify({'error': str(e)}), 500

def format_file_size(size_bytes):
    """Formata tamanho do arquivo em formato legível"""
    if size_bytes == 0:
        return "0 B"
    size_names = ["B", "KB", "MB", "GB"]
    i = int(math.floor(math.log(size_bytes, 1024)))
    p = math.pow(1024, i)
    s = round(size_bytes / p, 2)
    return f"{s} {size_names[i]}"
