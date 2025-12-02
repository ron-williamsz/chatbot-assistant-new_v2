from celery import Celery
from docx import Document
from datetime import datetime
import assemblyai as aai
import os
import logging
import redis
import time

# Configuração de logs mais detalhada
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('transcrever')

# Configuração do Celery
# Detecta se estamos rodando em ambiente Docker ou direto na máquina
redis_host = os.environ.get("REDIS_HOST", "localhost")
broker_url = os.environ.get("CELERY_BROKER_URL", f"redis://{redis_host}:6379/0")
result_backend = os.environ.get("CELERY_RESULT_BACKEND", f"redis://{redis_host}:6379/0")

app = Celery('tasks')
app.conf.update(
    broker_url=broker_url,
    result_backend=result_backend,
    task_serializer='json',
    result_serializer='json',
    accept_content=['json'],
    task_track_started=True,
    worker_pool='solo',
    broker_connection_retry_on_startup=True
)

# Verificar conexão com o Redis
redis_port = int(os.environ.get("REDIS_PORT", "6379"))

def verificar_redis():
    try:
        r = redis.Redis(host=redis_host, port=redis_port, db=0)
        r.ping()
        logger.info(f"Conexão com Redis estabelecida com sucesso em {redis_host}:{redis_port}!")
        return True
    except Exception as e:
        logger.error(f"Erro ao conectar ao Redis em {redis_host}:{redis_port}: {e}")
        return False

# Configuração da API AssemblyAI
# Obtém a chave da API de variável de ambiente ou usa a chave padrão como fallback
aai.settings.api_key = os.environ.get("ASSEMBLYAI_API_KEY", "e19ccd45d3c944a3a0614a6baffe2804")
logger.info(f"Usando chave da API AssemblyAI: {aai.settings.api_key[:5]}...{aai.settings.api_key[-4:]} (tamanho: {len(aai.settings.api_key)})")

@app.task(bind=True)
def process_file(self, filepath, language='pt', speaker_labels=True):
    try:
        verificar_redis()
        logger.info(f"Iniciando processamento do arquivo: {filepath}")
        logger.info(f"Parâmetros: language={language}, speaker_labels={speaker_labels}")
        self.update_state(state='PROGRESS', meta={'status': 'Iniciando transcrição...'})

        # Verificar se o arquivo existe
        if not os.path.exists(filepath):
            logger.error(f"Arquivo não encontrado: {filepath}")
            raise FileNotFoundError(f"Arquivo não encontrado: {filepath}")

        self.update_state(state='PROGRESS', meta={'status': 'Enviando para AssemblyAI...'})
        
        # Verificar chave da API
        if not aai.settings.api_key or len(aai.settings.api_key) < 10:
            error_msg = "Chave da API AssemblyAI inválida ou não configurada. Verifique o arquivo docker-compose.yml e adicione sua chave válida."
            logger.error(error_msg)
            raise ValueError(error_msg)
            
        # Enviar arquivo para AssemblyAI com configuração personalizada
        try:
            if speaker_labels:
                # Configuração otimizada para a diarização (identificação de interlocutores)
                # Definimos o número esperado de interlocutores para 5
                # De acordo com a documentação, os melhores resultados são obtidos quando:
                # 1. Cada interlocutor fala por pelo menos 30 segundos no total
                # 2. Não há muito ruído de fundo ou eco
                # 3. Há pouca sobreposição de fala entre interlocutores
                
                # Importante: A API suporta no máximo 10 interlocutores
                
                # Número estimado de interlocutores (ajuste conforme necessário)
                num_speakers = int(os.environ.get("SPEAKERS_EXPECTED", "5"))
                logger.info(f"Configurando diarização para {num_speakers} interlocutores esperados")
                
                config = aai.TranscriptionConfig(
                    language_code=language,      # Código do idioma configurável
                    punctuate=True,              # Adicionar pontuação à transcrição
                    format_text=True,            # Formatar texto automaticamente
                    speaker_labels=True,         # Ativar identificação de interlocutores
                    speakers_expected=num_speakers  # Define o número esperado de interlocutores
                )
            else:
                # Configuração sem diarização
                config = aai.TranscriptionConfig(
                    language_code=language,
                    punctuate=True,
                    format_text=True,
                    speaker_labels=False
                )

            logger.info(f"Iniciando transcrição com configuração: language={language}, speaker_labels={speaker_labels}")
            if speaker_labels:
                logger.info(f"Diarização ativada com {num_speakers} speakers esperados")
            
            transcriber = aai.Transcriber()

            # Log de início da transcrição
            start_time = time.time()
            
            # Executar a transcrição
            transcript = transcriber.transcribe(
                filepath,
                config=config
            )
            
            # Log de término da transcrição
            logger.info(f"Transcrição concluída em {time.time() - start_time:.2f} segundos")
            logger.info(f"Status da transcrição: {transcript.status}")
            
            # Verificar se a transcrição foi bem-sucedida
            if not transcript or transcript.status != "completed":
                logger.error("Transcrição não foi concluída com sucesso")
                raise ValueError("A API não conseguiu completar a transcrição. Verifique o arquivo de áudio.")
            
            # Verificar se temos texto transcrito
            if not hasattr(transcript, 'text') or not transcript.text:
                logger.error("Transcrição vazia retornada pela API")
                raise ValueError("A API não retornou texto transcrito. Verifique o arquivo de áudio.")
                
        except Exception as api_error:
            logger.error(f"Erro na API AssemblyAI: {str(api_error)}")
            raise ValueError(f"Falha na API de transcrição: {str(api_error)}. Verifique se a chave da API está correta em docker-compose.yml.")
            
        self.update_state(state='PROGRESS', meta={'status': 'Criando documento Word...'})
        
        # Criar documento Word
        doc = Document()
        doc.add_heading('Transcrição de Áudio', 0)
        
        current_time = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        doc.add_paragraph(f'Data de geração: {current_time}')
        doc.add_paragraph(f'Arquivo original: {os.path.basename(filepath)}')
        doc.add_paragraph('')
        
        # Verificar se temos diarização (utterances)
        has_utterances = (speaker_labels and 
                         hasattr(transcript, 'utterances') and 
                         transcript.utterances and 
                         len(transcript.utterances) > 0)
        
        logger.info(f"Verificando diarização: speaker_labels={speaker_labels}, has_utterances={has_utterances}")
        
        if has_utterances:
            # Log para verificar quantos interlocutores foram identificados
            speakers = set(utterance.speaker for utterance in transcript.utterances)
            logger.info(f"Interlocutores identificados: {', '.join(speakers)} (total: {len(speakers)})")
            
            # Analisar duração da fala de cada interlocutor (útil para diagnóstico)
            speaker_durations = {}
            for utterance in transcript.utterances:
                speaker = utterance.speaker
                duration = (utterance.end - utterance.start) / 1000  # conversão para segundos
                
                if speaker not in speaker_durations:
                    speaker_durations[speaker] = 0
                    
                speaker_durations[speaker] += duration
            
            # Registrar a duração da fala de cada interlocutor
            for speaker, duration in speaker_durations.items():
                logger.info(f"Interlocutor {speaker} falou por {duration:.2f} segundos no total")
                # A documentação recomenda pelo menos 30 segundos por interlocutor
                if duration < 30:
                    logger.warning(f"Interlocutor {speaker} falou por menos de 30 segundos, o que pode afetar a precisão da identificação")
            
            # Adicionar informações sobre os interlocutores identificados
            doc.add_heading('Informações sobre Interlocutores', level=1)
            doc.add_paragraph(f'Interlocutores identificados: {len(speakers)}')
            for speaker, duration in speaker_durations.items():
                doc.add_paragraph(f"Interlocutor {speaker}: {duration:.2f} segundos de fala")
            
            doc.add_paragraph('')
            doc.add_heading('Transcrição com Identificação de Falantes', level=1)
            
            # Adicionar o texto transcrito separado por locutor com timestamps
            for utterance in transcript.utterances:
                # Formatação mais clara com timestamp
                timestamp = f"{utterance.start/1000:.1f}s - {utterance.end/1000:.1f}s"
                text = f"Interlocutor {utterance.speaker} [{timestamp}]: {utterance.text}"
                doc.add_paragraph(text)
                
            result_speakers = len(speakers)
            
        else:
            # Sem diarização ou diarização falhou
            if speaker_labels:
                logger.warning("Diarização foi solicitada mas não há utterances disponíveis")
                doc.add_paragraph("AVISO: Identificação de falantes foi solicitada mas não foi possível identificar múltiplos falantes nesta transcrição.")
                doc.add_paragraph("")
            
            doc.add_heading('Transcrição Completa', level=1)
            doc.add_paragraph(transcript.text)
            result_speakers = 0 if speaker_labels else None
        
        # Gerar nomes únicos para os arquivos
        current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"transcricao_{current_time}.docx"
        
        docx_filepath = os.path.join('processed', docx_filename)

        # Criar diretório processed se não existir
        os.makedirs('processed', exist_ok=True)

        # Salvar DOCX
        doc.save(docx_filepath)
        logger.info(f"Arquivo DOCX salvo: {docx_filepath}")
        
        # Limpar arquivo de entrada
        if os.path.exists(filepath):
            os.remove(filepath)
            logger.info(f"Arquivo temporário removido: {filepath}")
        
        # Resultado final
        self.update_state(state='PROGRESS', meta={'status': 'Finalizado!'})
        
        result = {
            'status': 'SUCCESS',
            'docx': docx_filename,
            'message': 'Transcrição concluída com sucesso!'
        }
        
        if result_speakers is not None:
            result['speakers_identified'] = result_speakers
            
        return result
        
    except Exception as e:
        logger.error(f"Erro no processamento: {str(e)}")
        self.update_state(
            state='FAILURE',
            meta={
                'exc_type': type(e).__name__,
                'exc_message': str(e),
                'error': str(e)
            }
        )
        raise 

# Verificar conexão com o Redis ao inicializar
if __name__ == "__main__":
    verificar_redis() 